"""
제안서 AI 어시스턴트
3-Step RAG: 니즈 분석 → 모듈 검색 → 커리큘럼 조합 → 양식 출력
"""

import streamlit as st
import json
import os
import re
import time
import io
from google import genai
import chromadb
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============ 설정 ============
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MODULE_DB_PATH = os.path.join(SCRIPT_DIR, "module_db")
LEGACY_DB_PATH = os.path.join(SCRIPT_DIR, "chroma_db")  # 구버전 폴백용

# ── API 키 풀 로드 (GEMINI_API_KEY_1 ~ _9, 없으면 GEMINI_API_KEY 단일 키) ──
def _load_api_keys() -> list:
    keys = []
    for i in range(1, 11):
        k = st.secrets.get(f"GEMINI_API_KEY_{i}") or os.getenv(f"GEMINI_API_KEY_{i}")
        if k:
            keys.append(k)
    if not keys:
        single = st.secrets.get("GEMINI_API_KEY") or os.getenv("GEMINI_API_KEY")
        if single:
            keys.append(single)
    return keys

_API_KEYS = _load_api_keys()

if not _API_KEYS:
    st.error("❌ GEMINI_API_KEY가 설정되지 않았습니다. Streamlit Cloud → Settings → Secrets에 추가하세요.")
    st.stop()

# 현재 사용 중인 키 인덱스 (세션 내 유지)
if "api_key_index" not in st.session_state:
    st.session_state.api_key_index = 0

def _get_client() -> genai.Client:
    """현재 활성 키로 클라이언트 반환"""
    idx = st.session_state.get("api_key_index", 0) % len(_API_KEYS)
    return genai.Client(api_key=_API_KEYS[idx])

def _rotate_key() -> bool:
    """다음 키로 교체. 더 이상 키가 없으면 False 반환"""
    current = st.session_state.get("api_key_index", 0)
    if current + 1 < len(_API_KEYS):
        st.session_state.api_key_index = current + 1
        return True
    return False

# 임베딩용 클라이언트 (첫 번째 키 고정 사용)
client_genai = genai.Client(api_key=_API_KEYS[0])

MODEL_LITE   = "gemini-3.1-flash-lite"   # Step 1~3: 니즈 분석, 모듈 검색
MODEL_HEAVY  = "gemini-2.5-flash-lite"            # Step 4~6: 커리큘럼 설계, AI 검수, 재작성

def _generate(prompt: str, json_mode: bool = False, heavy: bool = False) -> str:
    """
    API 호출 공통 헬퍼.
    - heavy=False(기본): MODEL_LITE(3.1) 우선, 실패 시 MODEL_HEAVY(2.5) 폴백
    - heavy=True       : MODEL_HEAVY(2.5) 우선, 실패 시 MODEL_LITE(3.1) 폴백
    - 429(할당량 초과) → 다음 API 키로 자동 교체 후 재시도
    - 503 / UNAVAILABLE → 최대 3회 재시도 (지수 대기)
    """
    cfg = {"response_mime_type": "application/json"} if json_mode else {}
    last_exc = None
    model_order = [MODEL_HEAVY, MODEL_LITE] if heavy else [MODEL_LITE, MODEL_HEAVY]

    for model in model_order:
        for attempt in range(3):
            try:
                client = _get_client()
                resp = client.models.generate_content(
                    model=model, contents=prompt,
                    config=cfg if cfg else None,
                )
                return resp.text
            except Exception as e:
                last_exc = e
                err = str(e)
                if "429" in err or "RESOURCE_EXHAUSTED" in err:
                    # 할당량 초과 → 다음 키로 교체
                    rotated = _rotate_key()
                    if rotated:
                        continue  # 새 키로 즉시 재시도
                    else:
                        break  # 키 소진, 폴백 모델로
                elif "503" in err or "UNAVAILABLE" in err:
                    if attempt < 2:
                        wait = 30 * (attempt + 1)
                        m = re.search(r'retry[^\d]*(\d+(?:\.\d+)?)\s*s', err, re.IGNORECASE)
                        if m:
                            wait = int(float(m.group(1))) + 5
                        time.sleep(wait)
                    else:
                        break  # 폴백 모델로
                else:
                    break  # 다른 오류는 즉시 폴백 모델로

    raise Exception(f"API 호출 실패 (모든 키·모델 시도): {last_exc}")

# ============ [P0-A] Interactive Needs Gathering ============
# 필수 정보 체크리스트
REQUIRED_INFO = {
    "교육인원": {
        "keywords": ["명", "명정도", "人", "인원", "참석자", "people", "persons"],
        "pattern": r'(\d+)\s*명',
        "required": True,
        "description": "몇 명을 대상으로 하는가?"
    },
    "pain_point": {
        "keywords": ["어려움", "문제", "과제", "부족", "고민", "이슈", "어려워", "힘들", "도전"],
        "pattern": None,
        "required": True,
        "description": "가장 시급한 문제 1가지는?"
    },
    "실습비중": {
        "keywords": ["강의", "실습", "혼합", "비중", "실제", "연습", "실행", "워크숍", "토의"],
        "pattern": r'(강의|실습|혼합|워크숍|토의)',
        "required": True,
        "description": "강의형 vs 혼합형 vs 실습형?"
    },
    "기존시도": {
        "keywords": ["해봤", "시도", "경험", "했어", "작년", "지난해", "전에", "이전"],
        "pattern": None,
        "required": False,
        "description": "기존에 해본 교육이 있나?"
    }
}

def check_info_completeness(full_text: str) -> dict:
    """
    누적된 대화에서 각 필수 정보가 포함되었는지 확인
    Returns: {정보명: bool}
    """
    results = {}

    for info_key, info_spec in REQUIRED_INFO.items():
        # 키워드 기반 검사 (대소문자 무시)
        text_lower = full_text.lower()
        found_by_keyword = any(kw in text_lower for kw in info_spec["keywords"])

        # 정규식 기반 검사
        found_by_pattern = False
        if info_spec["pattern"]:
            found_by_pattern = bool(re.search(info_spec["pattern"], full_text))

        # 둘 중 하나라도 매치되면 포함된 것으로 간주
        results[info_key] = found_by_keyword or found_by_pattern

    return results

def is_needs_complete(full_text: str) -> bool:
    """
    필수 정보 3개 이상이 확보되었는지 판단
    Returns: bool
    """
    completeness = check_info_completeness(full_text)
    required_items = [k for k, v in REQUIRED_INFO.items() if v["required"]]
    found_count = sum(1 for k in required_items if completeness.get(k, False))

    # 필수 정보의 75% 이상 확보되면 검색 진행 가능
    threshold = int(len(required_items) * 0.75)
    return found_count >= threshold

def generate_follow_up_questions(initial_query: str, industry: str, target: str,
                                 conversation_history: list = None) -> str:
    """
    사용자의 초기 입력 + 현재까지 대화 내용을 분석하여
    부족한 정보를 파악하고 3~4개의 구체적인 질문 생성
    """

    # 누적된 대화 텍스트화
    accumulated_text = initial_query
    if conversation_history:
        accumulated_text += " " + " ".join([msg.get("content", "") for msg in conversation_history if msg.get("role") == "user"])

    # 현재 정보 완성도 확인
    completeness = check_info_completeness(accumulated_text)

    # 부족한 정보 식별
    missing_info = [key for key, complete in completeness.items()
                    if not complete and REQUIRED_INFO[key]["required"]]

    if not missing_info:
        # 모든 필수 정보가 확보된 경우
        return "완벽합니다! 충분한 정보가 수집되었습니다. '검색 진행' 버튼을 클릭해주세요."

    # 부족한 정보에 대한 질문 생성 (최대 3개)
    missing_questions_text = "\n".join([f"- {REQUIRED_INFO[key]['description']}" for key in missing_info[:3]])

    prompt = f"""당신은 경험 많은 HRD 컨설턴트입니다.

[초기 요청]
{initial_query}

[산업군]: {industry}
[교육 대상]: {target}

[이전 대화]
{chr(10).join([f"- 고객: {msg['content']}" for msg in conversation_history if msg.get('role') == 'user']) if conversation_history else "없음"}

[여전히 부족한 정보]
{missing_questions_text}

위 부족한 정보를 자연스럽게 얻기 위해 3~4개의 간단하고 구체적인 질문을 던져주세요.

예시 포맷:
"좋습니다! 몇 가지 더 확인하고 싶습니다:
1. [질문1]
2. [질문2]
3. [질문3]"

주의:
- 이미 언급된 정보는 다시 묻지 마세요
- 개방형 질문으로 자연스럽게
- 너무 길지 않게 (3~4개 질문만)
"""

    try:
        return _generate(prompt, heavy=True)
    except Exception:
        return f"질문 생성 중 오류가 발생했습니다. 직접 입력해주세요.\n\n부족한 정보: {missing_questions_text}"


# ============ [답변 2] 마크다운 → DOCX 변환 ============
def markdown_to_docx(markdown_text: str) -> bytes:
    """
    마크다운을 Word(.docx) 파일로 변환
    H1(#) → 제목1, H2(##) → 제목2, H3(###) → 제목3
    **bold** → 굵은 글씨, 표(|) → 워드 테이블, * 항목 → 글머리 기호
    """
    doc = Document()

    # 기본 폰트 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    def add_run_with_bold(para, text):
        """**text** 패턴을 파싱해서 굵은 글씨 처리"""
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                para.add_run(part)

    lines = markdown_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # 헤딩
        if line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)

        # 구분선
        elif line.strip() in ('---', '***', '___'):
            doc.add_paragraph('─' * 40)

        # 표 감지 (|로 시작하는 줄)
        elif line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            # 구분선 행 제거 (|---|---|)
            data_lines = [l for l in table_lines if not re.match(r'^\s*\|[\s\-:]+\|', l)]
            if not data_lines:
                continue

            # 셀 파싱
            rows = []
            for tl in data_lines:
                cells = [c.strip() for c in tl.strip().strip('|').split('|')]
                rows.append(cells)

            if rows:
                n_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=n_cols)
                table.style = 'Table Grid'
                for r_idx, row_data in enumerate(rows):
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < n_cols:
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                            if r_idx == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
            continue

        # 글머리 기호 (* 또는 -)
        elif re.match(r'^[\*\-] ', line):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_run_with_bold(p, text)

        # 들여쓰기 글머리 (  * 또는   -)
        elif re.match(r'^\s{2,}[\*\-] ', line):
            text = re.sub(r'^\s+[\*\-] ', '', line).strip()
            p = doc.add_paragraph(style='List Bullet 2')
            add_run_with_bold(p, text)

        # 빈 줄
        elif line.strip() == '':
            doc.add_paragraph('')

        # 일반 텍스트
        else:
            p = doc.add_paragraph()
            add_run_with_bold(p, line.strip())

        i += 1

    # bytes로 반환
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============ DB 로드 ============
@st.cache_resource
def load_module_db():
    """모듈 단위 ChromaDB 로드"""
    if os.path.exists(MODULE_DB_PATH):
        try:
            c = chromadb.PersistentClient(path=MODULE_DB_PATH)
            col = c.get_collection("modules")
            return col, "module"
        except Exception as e:
            print(f"[경고] module_db 로드 실패: {e}")
            st.warning(f"⚠️ module_db 로드 실패 ({e}). 구버전 DB로 폴백합니다.")
    else:
        print(f"[경고] module_db 디렉토리가 없습니다: {MODULE_DB_PATH}")
        st.warning(
            "⚠️ module_db가 없습니다. 구버전 DB로 폴백합니다.\n\n"
            "DB를 새로 구축하려면 터미널에서 아래 명령을 실행하세요:\n"
            "`python step4_build_module_db.py`"
        )

    # 폴백: 구버전 제안서 단위 DB
    if not os.path.exists(LEGACY_DB_PATH):
        raise FileNotFoundError(
            f"module_db도, 구버전 chroma_db도 존재하지 않습니다.\n"
            f"먼저 `python step4_build_module_db.py`를 실행하여 DB를 구축하세요."
        )
    c = chromadb.PersistentClient(path=LEGACY_DB_PATH)
    col = c.get_collection("proposals")
    return col, "legacy"


# ============ Step 1: 니즈 분석 ============
def analyze_needs(query, industry, target, duration):
    """
    Gemini API로 자연어 니즈 → 구조화된 JSON 변환
    """
    duration_h = int(duration)

    prompt = f"""당신은 10년 경력의 시니어 HRD 컨설턴트입니다.
고객이 입력한 교육 니즈를 분석하여 아래 JSON 형식으로 정확하게 변환하세요.

[고객 입력]
니즈: {query}
UI 산업군(참고용): {industry}
UI 교육 대상(참고용): {target}
교육 시간: {duration_h}H

[출력 규칙]
- target: 교육 대상 직급. 니즈 텍스트에서 구체적으로 언급된 경우 그 내용을 우선 반영. 아니면 UI 값 사용.
  (예시: "팀장/리더급", "신입사원", "실무진", "중간관리자", "임원", "전직급")
- industry: 산업군. 니즈 텍스트에서 구체적으로 언급된 경우 그 내용을 우선 반영. 아니면 UI 값 사용.
  (예시: "전산업", "제조", "금융/은행", "공공기관", "IT/테크", "건설/부동산", "유통/서비스")
- core_keywords: 교육에서 반드시 다뤄야 할 핵심 역량/주제 키워드 3~5개 (산업군·직급·시간 등 메타 정보는 제외하고 실제 교육 내용 키워드만 추출)
- pain_point: 현재 조직/구성원의 핵심 문제점 1~2문장
- expected_behavior: 교육 후 기대되는 구체적 행동 변화 1~2문장
- learning_level: beginner / intermediate / advanced 중 선택

반드시 아래 JSON 형식으로만 응답하세요. 다른 텍스트는 절대 포함하지 마세요.

{{
  "target": "니즈에서 추론한 교육 대상",
  "industry": "니즈에서 추론한 산업군",
  "duration_hours": {duration_h},
  "core_keywords": ["키워드1", "키워드2", "키워드3"],
  "pain_point": "현재 문제점",
  "expected_behavior": "기대 행동 변화",
  "learning_level": "intermediate",
  "preferred_style": "실습형"
}}"""

    try:
        return json.loads(_generate(prompt, json_mode=True, heavy=True))
    except Exception as e:
        # JSON 파싱 실패 시 기본값 반환
        return {
            "target": target,
            "industry": industry,
            "duration_hours": duration_h,
            "core_keywords": query.split()[:5],
            "pain_point": query,
            "expected_behavior": "역량 향상",
            "learning_level": "intermediate",
            "preferred_style": "실습형"
        }


# ============ Step 2: 모듈 검색 ============
# ============ [P0-B] Hard-Binding RAG Context ============

_RRF_K = 60  # RRF 상수

def _embed_batch(queries: list) -> list:
    """여러 텍스트를 1회 배치 API 호출로 임베딩 (429 재시도 포함)"""
    import re as _re
    for attempt in range(5):
        try:
            result = client_genai.models.embed_content(
                model="gemini-embedding-001",
                contents=queries
            )
            return [e.values for e in result.embeddings]
        except Exception as e:
            if "429" in str(e) and attempt < 4:
                m = _re.search(r'retry[^\d]*(\d+(?:\.\d+)?)\s*s', str(e), _re.IGNORECASE)
                wait = int(float(m.group(1))) + 5 if m else 30 * (attempt + 1)
                print(f"[임베딩 rate limit] {wait}초 후 재시도 (시도 {attempt+1}/5)")
                time.sleep(wait)
            else:
                raise e
    raise Exception("임베딩 생성 실패: 최대 재시도 초과")


def _generate_search_queries(needs_json: dict) -> list:
    """
    Gemini로 모듈 검색에 최적화된 쿼리 3개 생성.
    - 주제 특화 동의어/관련어 확장
    - 하드코딩 쿼리 대비 훨씬 높은 검색 적합도
    - 실패 시 기본 쿼리로 폴백
    """
    kw       = needs_json.get("core_keywords", [])
    pain     = needs_json.get("pain_point", "")
    behavior = needs_json.get("expected_behavior", "")
    target   = needs_json.get("target", "")
    industry = needs_json.get("industry", "")

    prompt = f"""당신은 HRD 교육 모듈 DB 검색 전문가입니다.
아래 교육 니즈를 보고, 벡터 DB 검색에 최적화된 검색 쿼리 3개를 생성하세요.

교육 니즈:
- 교육 대상: {target}
- 산업군: {industry}
- 핵심 키워드: {', '.join(kw)}
- 문제점: {pain}
- 기대 행동 변화: {behavior}

규칙:
1. 쿼리 1: 핵심 주제와 관련 동의어·유사어를 모두 포함. 교육 대상({target})과 산업군({industry})에 맞는 맥락 용어도 추가 (예: 팀장 대상 금융권 → "팀장 리더십 금융 지점장 성과관리")
2. 쿼리 2: 교육 내용 관점 — {target}이 실제로 배울 법한 구체적 학습 내용·스킬. {industry} 업종 특성 반영
3. 쿼리 3: 현업 적용 관점 — 교육 후 {target}이 {industry} 현업에서 실제로 쓰는 상황·행동

반드시 JSON 배열로만 응답: ["쿼리1", "쿼리2", "쿼리3"]"""

    try:
        queries = json.loads(_generate(prompt, json_mode=True, heavy=True))
        if isinstance(queries, list) and len(queries) >= 3:
            print(f"[검색 쿼리 생성] {queries}")
            return [str(q) for q in queries[:3]]
    except Exception as e:
        print(f"[검색 쿼리 생성 실패, 기본 쿼리 사용] {e}")

    # 폴백: 기본 쿼리 (target/industry 포함)
    return [
        f"{target} {industry} " + " ".join(kw),
        f"{target} {pain} {behavior}",
        f"{industry} " + " ".join(kw) + f" {pain}",
    ]


def search_modules_detailed(collection, needs_json, db_type):
    """
    멀티쿼리 검색 + RRF 병합: 상위 20개 반환
    - LLM이 주제 최적화 쿼리 3개 생성 (동의어 확장, 주제 집중)
    - 배치 임베딩 1회 API 호출
    - RRF로 공통 상위 모듈 우선 정렬
    """
    queries = _generate_search_queries(needs_json)

    # 3개 쿼리를 배치로 임베딩 (API 호출 1회)
    embeddings = _embed_batch(queries)

    # RRF 점수 계산
    rrf_scores = {}  # doc_id → RRF score
    id_to_meta = {}  # doc_id → metadata
    id_to_sim  = {}  # doc_id → 최고 유사도 (UI 표시용)

    n_results = min(100, collection.count())
    for embedding in embeddings:
        raw = collection.query(
            query_embeddings=[embedding],
            n_results=n_results
        )
        for rank_0idx, (meta, dist, doc_id) in enumerate(zip(
            raw["metadatas"][0],
            raw["distances"][0],
            raw["ids"][0]
        )):
            rank = rank_0idx + 1
            rrf_scores[doc_id] = rrf_scores.get(doc_id, 0.0) + 1.0 / (_RRF_K + rank)
            similarity = round((1 - dist) * 100, 1)
            if doc_id not in id_to_meta or id_to_sim[doc_id] < similarity:
                id_to_meta[doc_id] = meta
                id_to_sim[doc_id]  = similarity

    # RRF 내림차순 정렬, 상위 20개
    top_ids = sorted(rrf_scores, key=lambda x: rrf_scores[x], reverse=True)[:20]

    retrieved_modules = []
    if db_type == "module":
        for idx, doc_id in enumerate(top_ids):
            sim = id_to_sim[doc_id]
            if sim < 70:  # 유사도 70% 미만 제외
                continue
            meta = id_to_meta[doc_id]
            retrieved_modules.append({
                "rank": len(retrieved_modules) + 1,
                "similarity_percent": sim,
                "모듈명":       meta.get("모듈명", ""),
                "과정명":       meta.get("과정명", ""),
                "교육목표":     meta.get("교육목표", ""),
                "내용_원문":    meta.get("내용_원문", ""),
                "권장시간":     meta.get("권장시간", ""),
                "교육방식":     meta.get("교육방식", ""),
                "모듈성격":     meta.get("모듈성격", "core"),
            })
    else:
        # 구버전 폴백: 단일 쿼리
        raw = collection.query(query_embeddings=[embeddings[0]], n_results=min(20, collection.count()))
        for idx, (meta, dist) in enumerate(zip(raw["metadatas"][0], raw["distances"][0])):
            sim = round((1 - dist) * 100, 1)
            if sim < 70:
                continue
            retrieved_modules.append({
                "rank": len(retrieved_modules) + 1,
                "similarity_percent": sim,
                "과정명":       meta.get("과정명", ""),
                "내용_원문":    meta.get("내용_원문", ""),
                "권장시간":     meta.get("권장시간", ""),
            })

    # ── 정렬: similarity_percent 내림차순 ──
    retrieved_modules.sort(key=lambda x: x["similarity_percent"], reverse=True)

    # ── 비즈니스 맥락 분석: 추천 타겟 / 관련 산업 ──
    if retrieved_modules and db_type == "module":
        items = ""
        for i, m in enumerate(retrieved_modules):
            course  = m.get("과정명", "")
            name    = m.get("모듈명", "")
            content = m.get("내용_원문", "")[:200]
            items  += f"{i}. 과정명: {course} | 모듈명: {name} | 내용: {content}\n"

        ctx_prompt = f"""아래 교육 모듈들을 보고 각 모듈의 추천 타겟 직급과 관련 산업을 추론하세요.

모듈 목록:
{items}

규칙:
- 추천타겟: 아래 중 가장 적합한 1~2개만 선택
  선택지: 전직급, 신입/주니어, 실무자, 팀장/리더급, 중간관리자, 임원, C-Level
- 관련산업: 아래 중 가장 적합한 1~2개만 선택
  선택지: 전산업, 금융/은행, 제조, 공공기관, IT/테크, 건설/부동산, 에너지/유틸리티, 유통/서비스, 의료/헬스케어

반드시 아래 JSON 형식으로만 응답:
{{"0": {{"추천타겟": ["팀장/리더급"], "관련산업": ["금융/은행"]}}, "1": {{"추천타겟": ["전직급"], "관련산업": ["전산업"]}}}}"""

        try:
            ctx_result = json.loads(_generate(ctx_prompt, json_mode=True, heavy=True))
            for i, m in enumerate(retrieved_modules):
                if isinstance(ctx_result, list):
                    ctx = ctx_result[i] if i < len(ctx_result) else {}
                else:
                    ctx = ctx_result.get(str(i), {})
                m["추천타겟"] = ctx.get("추천타겟", [])
                m["관련산업"] = ctx.get("관련산업", [])
        except Exception as e:
            print(f"[맥락 분석 실패] {e}")
            for m in retrieved_modules:
                m["추천타겟"] = []
                m["관련산업"] = []

    retrieved_json = json.dumps(retrieved_modules, ensure_ascii=False, indent=2)
    return retrieved_modules, retrieved_json


def group_modules_by_type(retrieved_modules, db_type):
    """
    retrieved_modules 리스트를 모듈성격별로 그룹핑: intro / core / apply
    """
    groups = {"intro": [], "core": [], "apply": []}

    for m in retrieved_modules:
        module_type = m.get("모듈성격", "core") if db_type == "module" else "core"
        if module_type not in groups:
            module_type = "core"
        groups[module_type].append({
            "meta": {k: v for k, v in m.items() if k not in ("rank", "similarity_percent", "모듈성격")},
            "similarity": m.get("similarity_percent", 0),
        })

    return groups


# ============ Step 3: 커리큘럼 조합 ============
# ============ [P1-1] 활동 소요 시간 검증 ============
def validate_curriculum_timing(curriculum_text: str, total_hours: int) -> dict:
    """
    생성된 커리큘럼에서 활동별 소요 시간을 추출하고 검증
    Returns: {valid: bool, total_minutes: int, warnings: list, details: str}
    """
    total_minutes = total_hours * 60
    extracted_activities = []

    # 표 형식의 커리큘럼에서 시간 추출 (정규식)
    table_pattern = r'\|\s*(.+?)\s*\|\s*(\d+(?:\.\d+)?)\s*[HhWw]?\s*\|'
    matches = re.findall(table_pattern, curriculum_text)

    total_allocated = 0
    for module_name, hours_str in matches:
        try:
            hours = float(hours_str)
            minutes = hours * 60
            total_allocated += minutes
            extracted_activities.append({
                "모듈": module_name.strip()[:40],
                "시간": hours,
                "분": minutes
            })
        except ValueError:
            pass

    # 검증
    is_valid = total_allocated <= total_minutes
    variance = total_minutes - total_allocated

    details = f"배분된 시간: {total_allocated}분 / 전체: {total_minutes}분 (여유: {variance}분)"

    warnings = []
    if variance < 0:
        warnings.append(f"⚠️ 교육 시간 초과: {abs(variance)}분 과다")
    elif variance < 30:
        warnings.append(f"⚠️ 여유 시간 부족: {variance}분만 남음 (피드백/휴식 시간 추가 권장)")

    return {
        "valid": is_valid,
        "total_minutes_allocated": total_allocated,
        "total_minutes_available": total_minutes,
        "variance_minutes": variance,
        "activities_count": len(extracted_activities),
        "warnings": warnings,
        "details": details,
        "activities": extracted_activities
    }


def validate_proposal_quality(text: str) -> dict:
    """
    [P1] AI Self-Correction Validator
    생성된 제안서가 최소 품질 기준을 충족하는지 검증
    Returns: {passed: bool, failures: list, module_count: int, interactions_found: list}
    """
    failures = []

    # 검증 1: 모듈 헤더(### N.) 개수 (3개 이상)
    module_count = len(re.findall(r'^###\s+\d+[\.\s]', text, re.MULTILINE))
    if module_count < 3:
        failures.append(f"커리큘럼 모듈이 {module_count}개로 3개 미만입니다. (최소 3개 필요)")

    # 검증 2: 세부 항목 bullet 개수 (전체 모듈 합산 기준 — 모듈당 평균 3개 이상)
    bullet_count = len(re.findall(r'(?:^|\n)\s*[·\-\*]\s+\S', text))
    if module_count > 0 and bullet_count < module_count * 3:
        failures.append(
            f"모듈당 세부 항목이 부족합니다. "
            f"(현재 전체 {bullet_count}개 / 최소 {module_count * 3}개 필요)"
        )

    # 검증 3: 상호작용 활동 키워드 포함 여부
    interaction_keywords = ['토의', '실습', '롤플레잉', '롤플레이', '워크샵', '사례분석', '케이스스터디']
    found_interactions = [kw for kw in interaction_keywords if kw in text]
    if not found_interactions:
        failures.append(
            "상호작용형 활동 키워드(토의/실습/롤플레잉/워크샵/사례분석)가 전혀 포함되지 않았습니다."
        )

    return {
        "passed": len(failures) == 0,
        "failures": failures,
        "module_count": module_count,
        "interactions_found": found_interactions,
    }


def assemble_curriculum(needs_json, grouped_modules, duration, retrieved_modules_json=None, selected_modules=None, track="standard", advanced_context=None):
    """
    Gemini가 검색된 모듈을 조합하여 최적 커리큘럼 구성
    [P0-B] Hard-Binding RAG: Retrieved modules를 JSON으로 강제 주입
    [P1-1] CoT Time Validator: AI가 활동별 소요 시간을 계산하고 자동 조정
    [PM] selected_modules: 사용자가 선택한 모듈 (우선 포함 필수)
    """
    total_h = int(duration)
    total_minutes = total_h * 60

    # 각 그룹에서 상위 모듈 선별 (화면 표시용 - 프롬프트에는 사용 안 함)
    def format_modules(modules, limit=5):
        out = []
        for m in modules[:limit]:
            meta = m["meta"]
            name = meta.get("모듈명", "")
            topics = meta.get("교육목표", "")
            content = meta.get("내용_원문", "")[:150]
            course = meta.get("과정명", "")
            sim = m["similarity"]
            out.append(f"- [{course}] {name}\n  주제: {topics}\n  내용: {content}\n  유사도: {sim}%")
        return "\n".join(out) if out else "없음"

    intro_list = format_modules(grouped_modules.get("intro", []))
    core_list = format_modules(grouped_modules.get("core", []), 8)
    apply_list = format_modules(grouped_modules.get("apply", []))

    # [P0-B] retrieved_modules_json이 없으면 이전 방식으로 폴백
    if not retrieved_modules_json:
        retrieved_modules_json = json.dumps([
            {"모듈명": intro_list, "내용_원문": "도입 모듈"},
            {"모듈명": core_list, "내용_원문": "핵심 모듈"},
            {"모듈명": apply_list, "내용_원문": "현업 적용 모듈"}
        ], ensure_ascii=False, indent=2)

    # [PM] 사용자 선택 모듈 섹션 생성
    if selected_modules:
        selected_json_str = json.dumps(selected_modules, ensure_ascii=False, indent=2)
        selected_section = f"""
## 🎯 [사용자 선택 모듈] 반드시 우선 포함 (필수)

사용자가 직접 선택한 모듈입니다. 아래 모듈들은 커리큘럼에 **반드시** 포함되어야 합니다.

```json
{selected_json_str}
```

### 선택 모듈 우선 순위 규칙
1. 위 선택된 모듈들은 100% 커리큘럼에 포함합니다 (생략 불가).
2. 선택된 모듈의 "내용_원문"에서 구체적 활동을 추출하여 그대로 사용합니다.
3. 남은 시간은 아래 Retrieved Modules에서 가장 적합한 모듈로 채웁니다.
"""
    else:
        selected_section = """
## 🎯 모듈 선택 안내
사용자가 특별히 선택한 모듈이 없습니다. 아래 Retrieved Modules 목록에서 AI가 가장 적합한 모듈을 자유롭게 선택하세요.
"""

    prompt = f"""당신은 10년 경력의 시니어 HRD 컨설턴트입니다.
아래 고객 니즈와 검색된 교육 모듈을 바탕으로, 완성도 높은 교육 제안서를 작성하세요.

## 고객 니즈 분석 결과
- 교육 대상: {needs_json.get('target')}
- 산업군: {needs_json.get('industry')}
- 핵심 키워드: {', '.join(needs_json.get('core_keywords', []))}
- 현재 문제점: {needs_json.get('pain_point')}
- 기대 행동 변화: {needs_json.get('expected_behavior')}
- 교육 시간: {total_h}H (= {total_minutes}분)
{selected_section}

## ⭐ [P0-B] Retrieved Modules (검색된 실제 교육 모듈들 - 반드시 이것만 사용)

다음은 고객의 니즈와 가장 유사한 실제 교육 모듈들입니다.
당신은 이 모듈들의 세부 활동과 시간을 "레고 블록처럼" 조립하여 새로운 커리큘럼을 구성해야 합니다.
이 목록에 없는 새로운 활동을 창작하면 안 됩니다.

```json
{_slim_modules_json(retrieved_modules_json, content_limit=400) if retrieved_modules_json else "{}"}
```

## 검색된 교육 모듈 풀 (참고용 - 위 JSON이 정확함)

### [도입 모듈 후보]
{intro_list}

### [핵심 스킬 모듈 후보]
{core_list}

### [현업 적용 모듈 후보]
{apply_list}

## 활동별 소요 시간 참고 데이터 (Chain-of-Thought 계산용)
- 강의/설명: 1.0분/분량
- 진단/사전평가: 5분
- 롤플레잉: 10분
- 페어 활동: 10분
- 팀/전체 토의: 8분
- 실습 활동: 12분
- 사례 분석: 8분
- 케이스 스터디: 12분
- 시뮬레이션: 15분
- 피드백 세션: 5분
- 발표/공유: 8분
- Q&A: 5분

## 작성 지침 (필수)

### [포맷팅 절대 규칙]
- 마크다운 표 내부 '세부 내용' 컬럼에 **절대로 `<br>` 태그나 HTML을 사용하지 마십시오.**
- 줄바꿈이 필요하면 반드시 `\n- ` 형식의 마크다운 리스트로 작성하십시오.

### [P0-B] Hard-Binding RAG 규칙 (100% 강제)
1. **절대 금지**: 위 "Retrieved Modules" 목록에 없는 활동을 추가하지 마세요.
2. **레고 블록 조립**: 각 모듈의 "내용_원문"에서 구체적인 활동(진단, 롤플레이, 토의 등)을 추출하여 그대로 사용하세요.
3. **시간 참고**: 각 모듈의 "권장시간"은 참고 기준입니다. 풍부한 내용 제공을 위해 시간이 다소 초과되더라도 내용을 축소하거나 활동을 삭제하지 마세요.
4. **충분한 내용**: 각 모듈당 최소 5~8개의 풍부한 세부 항목을 원본 JSON의 "내용_원문"에서 가져와 표 형식으로 구성하세요.

### [P0-2] RAG 강제 인용 규칙 (80% 이상)
1. **"검색된 교육 모듈 풀"의 모듈들의 구체적 활동을 반드시 80% 이상 커리큘럼에 인용하세요.**
   - 모듈 제목만 차용하지 말 것 (❌ "피드백 스킬")
   - 구체적 활동을 직접 명시할 것 (✅ "피어 피드백 진단지 작성 (30분)")
2. AI 창작은 최소화하고, 검색된 모듈의 검증된 콘텐츠를 우선 활용하세요.

### [핵심 품질 지침 — 반드시 준수]

**[지침 1] 내용 풍부성 (절대 축소 금지)**
- 절대 내용을 축소하거나 요약하지 마십시오.
- 컨설턴트가 취사선택할 수 있도록 각 모듈당 최소 5개 이상의 구체적인 세부 내용을 제안하십시오.
- 시간이 초과된다는 이유로 활동을 삭제하거나 모듈 수를 줄이는 것을 엄격히 금지합니다.

**[지침 2] 상호작용형 활동 필수 포함**
- 각 커리큘럼 모듈에는 이론/강의뿐만 아니라, 반드시 최소 1개 이상의 상호작용형 활동을 포함하십시오.
- 활동 유형: `[토의]`, `[실습]`, `[롤플레잉]`, `[워크샵]`, `[사례분석]` 중 선택
- 구체적인 실행 방법과 사례명을 명시하십시오.
  - 예: `[롤플레잉] 고객 클레임 3단계 대응 시나리오 — 3인 1조 역할 분담 후 피드백`
  - 예: `[토의] 우리 팀의 소통 장벽 Top 3 도출 — 포스트잇 브레인스토밍`

**[지침 3] 시간 표기 방식**
- 시간 컬럼에는 고정값 대신 **범위 형태**로 표기하세요. (예: `60~90분 (제안)`, `90~120분 (제안)`)
- 총 교육 시간 {total_h}H은 컨설턴트가 최종 조율하는 **참고 기준**입니다.
- 시간을 맞추기 위해 내용을 삭제하는 행위는 엄격히 금지합니다.

4. 반드시 아래 양식을 정확히 따라 작성하세요

---

# 맞춤 교육 제안서

## 📋 과정 개요
* **과정명:** (창의적이고 전문적인 과정명)
* **교육 대상:** {needs_json.get('target')}
* **교육 목적:** (개괄식으로 2개 작성)
* **교육 시간:** {total_h}H
* **교육 방식:** (강의형/실습형/워크샵형 중 선택)

---

## 📚 상세 커리큘럼

(아래 양식을 모듈마다 반복. 마크다운 표(table)는 절대 사용 금지)

### 1. [모듈 주제명] (60~90분 제안)
- 주요 학습 내용 요약 (1~2문장)
  - 세부 상세 내용 1
  - 세부 상세 내용 2
  - 세부 상세 내용 3
  - 세부 상세 내용 4
  - 세부 상세 내용 5
  - [실습] 구체적 활동명 및 실행 방법 (대괄호 필수)

### 2. [모듈 주제명] (60~90분 제안)
- 주요 학습 내용 요약
  - 세부 상세 내용 1~5
  - [토의] 구체적 토의 주제 및 진행 방식 (대괄호 필수)

### 3. [모듈 주제명] (60~90분 제안)
- 주요 학습 내용 요약
  - 세부 상세 내용 1~5
  - [롤플레잉] 시나리오명 및 역할 구성 (대괄호 필수)

(모듈은 교육 시간에 따라 자유롭게 추가. 표 사용 절대 금지)
"""

    # [Sprint 2-2] 고도화 모드: 조직 맥락 컨텍스트 + 추가 섹션 강제
    if advanced_context:
        adv_filled = {k: v for k, v in advanced_context.items() if v and str(v).strip()}
        if adv_filled:
            ctx_lines = "\n".join(f"- {k}: {v}" for k, v in adv_filled.items())
            prompt = f"""[조직 맞춤화 컨텍스트 — 반드시 커리큘럼에 반영하세요]
{ctx_lines}

""" + prompt

    # [P1] Self-Correction 루프: 최대 3회 시도 (첫 생성 + 최대 2회 재생성)
    MAX_QUALITY_ATTEMPTS = 3
    last_curriculum = None

    for quality_attempt in range(MAX_QUALITY_ATTEMPTS):
        # API 호출 (재시도 + 폴백 포함)
        try:
            curriculum = _generate(prompt, heavy=True)
        except Exception as e:
            raise Exception(f"API 호출에 실패했습니다: {e}")

        if curriculum is None:
            raise Exception("API 호출에 실패했습니다.")

        last_curriculum = curriculum

        # [P1] Self-Correction Validator — 품질 검증
        quality_result = validate_proposal_quality(curriculum)

        if quality_result["passed"]:
            # 품질 기준 통과 → 바로 반환
            break

        if quality_attempt < MAX_QUALITY_ATTEMPTS - 1:
            # 품질 기준 미달 → 백그라운드 재생성 (에러 노출 없음)
            failures_text = " | ".join(quality_result["failures"])
            # 프롬프트 끝에 실패 원인을 추가하여 재시도
            prompt = prompt + f"""

---
[자동 재생성 — 이전 결과 품질 기준 미달]
이전 생성 결과가 아래 기준을 충족하지 못했습니다. 반드시 수정하여 재생성하세요:
{chr(10).join(f'- {f}' for f in quality_result['failures'])}
특히 각 모듈에 [토의]/[실습]/[롤플레잉]/[워크샵]/[사례분석] 중 하나 이상을 반드시 포함하고,
모듈당 세부 항목을 5개 이상 작성하세요.
"""
        # 마지막 시도이거나 통과했으면 루프 종료

    # 시간 검증 (참고용 — 내부 QA용)
    timing_result = validate_curriculum_timing(last_curriculum, total_h)
    st.session_state.curriculum_timing = timing_result

    return last_curriculum, timing_result


# ============ 프롬프트용 모듈 JSON 경량화 (토큰 절약) ============
def _slim_modules_json(retrieved_modules_json: str, content_limit: int = 300) -> str:
    """내용_원문을 content_limit자로 잘라 프롬프트 토큰 수를 줄임."""
    try:
        mods = json.loads(retrieved_modules_json)
        for m in mods:
            if "내용_원문" in m and len(m["내용_원문"]) > content_limit:
                m["내용_원문"] = m["내용_원문"][:content_limit] + "…"
        return json.dumps(mods, ensure_ascii=False)
    except Exception:
        return retrieved_modules_json


# ============ A/B 완성 제안서 생성 (CoT) ============
def assemble_curriculum_ab(needs_json, retrieved_modules_json, duration, selected_modules=None):
    """
    CoT 방식으로 서로 다른 교육 철학의 A안·B안 완성 제안서를 한 번에 생성.
    반환: (cot_text, draft_a, draft_b)
    - cot_text : AI 설계 사고 과정
    - draft_a  : A안 완성 제안서 (과정 개요 + 상세 커리큘럼)
    - draft_b  : B안 완성 제안서 (과정 개요 + 상세 커리큘럼)
    """
    total_h = int(duration)
    target   = needs_json.get("target", "")
    industry = needs_json.get("industry", "")
    keywords = ", ".join(needs_json.get("core_keywords", []))
    pain     = needs_json.get("pain_point", "")
    behavior = needs_json.get("expected_behavior", "")

    # 토큰 절약: 내용_원문을 300자로 제한
    slim_modules_json = _slim_modules_json(retrieved_modules_json, content_limit=300)

    sel_note = ""
    if selected_modules:
        names = [m.get("모듈명", "") for m in selected_modules]
        sel_note = f"\n⚠️ 아래 모듈은 사용자가 선택한 필수 포함 모듈입니다: {', '.join(names)}\n"

    proposal_format = f"""# 맞춤 교육 제안서

## 📋 과정 개요
* **과정명:** (창의적이고 전문적인 과정명)
* **교육 대상:** {target}
* **교육 목적:** (개괄식으로 2개 작성)
* **교육 시간:** {total_h}H
* **교육 방식:** (강의형/실습형/워크샵형 중 선택)

---

## 📚 상세 커리큘럼

(아래 양식을 모듈마다 반복. 마크다운 표(table)는 절대 사용 금지)

### 1. [모듈 주제명] (60~90분 제안)
- 주요 학습 내용 요약 (1~2문장)
  - 세부 상세 내용 1
  - 세부 상세 내용 2
  - 세부 상세 내용 3
  - 세부 상세 내용 4
  - 세부 상세 내용 5
  - [실습/토의/롤플레잉] 구체적 활동명 및 실행 방법 (대괄호 필수)

### 2. [모듈 주제명] (60~90분 제안)
...

(모듈은 교육 시간에 따라 자유롭게 추가. 표 사용 절대 금지)"""

    prompt = f"""당신은 시니어 HRD 컨설턴트입니다.
아래 교육 니즈와 모듈을 분석하여, 서로 다른 교육 철학의 완성된 교육 제안서 A안과 B안을 작성하세요.

## 교육 니즈
- 교육 대상: {target}
- 산업군: {industry}
- 핵심 키워드: {keywords}
- 문제점: {pain}
- 기대 행동 변화: {behavior}
- 교육 시간: {total_h}H
{sel_note}
## 검색된 교육 모듈 (이 모듈만 사용)
```json
{slim_modules_json}
```

## 작성 규칙
- 두 안은 서로 다른 교육 철학·구성 순서를 가져야 함 (단순 모듈 순서 변경 X)
- 각 안에서 모듈은 검색된 목록에서만 선택
- 각 모듈당 최소 5개 이상의 구체적 세부 내용 포함
- 마크다운 표(table) 사용 절대 금지
- 각 모듈에 [토의]/[실습]/[롤플레잉]/[워크샵]/[사례분석] 중 하나 이상 포함
- 시간 표기는 범위 형태로 (예: 60~90분 제안)
- 반드시 아래 구분자를 정확히 사용

먼저 아래 설계 사고를 작성하세요:

## 💭 AI 설계 사고 과정
- **핵심 니즈 해석:** (이 교육에서 가장 중요한 것은?)
- **A안 방향:** (예: 이론→스킬→실습 단계 심화형)
- **B안 방향:** (예: 현업 문제 중심 경험학습형)
- **두 안의 차별점:** (A/B를 구분하는 핵심 철학 차이)

그 다음 아래 구분자 형식으로 두 안의 완성된 제안서를 작성하세요:

---DRAFT_A_START---
{proposal_format}
---DRAFT_A_END---

---DRAFT_B_START---
{proposal_format}
---DRAFT_B_END---
"""

    try:
        text = _generate(prompt, heavy=True)
    except Exception as e:
        raise Exception(f"A/B 제안서 생성 실패: {e}")

    # CoT 파싱
    cot_text = ""
    if "## 💭 AI 설계 사고 과정" in text:
        cot_start = text.index("## 💭 AI 설계 사고 과정")
        cot_end   = text.index("---DRAFT_A_START---") if "---DRAFT_A_START---" in text else len(text)
        cot_text  = text[cot_start:cot_end].strip()

    # A안 파싱
    draft_a = ""
    if "---DRAFT_A_START---" in text and "---DRAFT_A_END---" in text:
        a_start = text.index("---DRAFT_A_START---") + len("---DRAFT_A_START---")
        a_end   = text.index("---DRAFT_A_END---")
        draft_a = text[a_start:a_end].strip()

    # B안 파싱
    draft_b = ""
    if "---DRAFT_B_START---" in text and "---DRAFT_B_END---" in text:
        b_start = text.index("---DRAFT_B_START---") + len("---DRAFT_B_START---")
        b_end   = text.index("---DRAFT_B_END---")
        draft_b = text[b_start:b_end].strip()

    return cot_text, draft_a, draft_b


# ============ A/B 통합 최선 제안서 생성 ============
def combine_ab_proposals(proposal_a, proposal_b, user_opinion, needs_json, duration):
    """
    A안·B안 완성 제안서를 비교하여 최선 요소를 통합한 최종 제안서 생성.
    사용자 의견을 반영하여 AI가 조합.
    반환: (combined_proposal, timing_result)
    """
    total_h = int(duration)
    total_minutes = total_h * 60
    target   = needs_json.get("target", "")
    keywords = ", ".join(needs_json.get("core_keywords", []))
    pain     = needs_json.get("pain_point", "")
    behavior = needs_json.get("expected_behavior", "")

    user_opinion_section = ""
    if user_opinion and user_opinion.strip():
        user_opinion_section = f"""
## 💬 사용자 의견 (반드시 반영)
{user_opinion.strip()}
"""

    prompt = f"""당신은 10년 경력의 시니어 HRD 컨설턴트입니다.
아래 두 가지 교육 제안서(A안, B안)를 비교 분석하여, 각 안의 장점을 통합한 최종 완성 제안서를 작성하세요.

## 고객 니즈
- 교육 대상: {target}
- 핵심 키워드: {keywords}
- 현재 문제점: {pain}
- 기대 행동 변화: {behavior}
- 교육 시간: {total_h}H (= {total_minutes}분)
{user_opinion_section}
## A안 제안서
{proposal_a}

## B안 제안서
{proposal_b}

## 작성 지침

### [통합 원칙]
1. A안과 B안 각각의 강점을 분석하고, 더 우수한 요소를 선택하거나 조합하세요.
2. 사용자 의견이 있다면 최우선으로 반영하세요.
3. 두 안의 단순 합산이 아니라, 진정으로 최선의 커리큘럼을 설계하세요.
4. 중복되는 모듈은 더 풍부한 내용을 가진 쪽을 채택하세요.

### [포맷팅 규칙]
- 마크다운 표(table) 사용 절대 금지
- 각 모듈당 최소 5개 이상의 구체적 세부 내용 포함
- 각 모듈에 [토의]/[실습]/[롤플레잉]/[워크샵]/[사례분석] 중 하나 이상 포함
- 시간 표기는 범위 형태로 (예: 60~90분 제안)

반드시 아래 양식으로 최종 제안서를 작성하세요:

---

# 맞춤 교육 제안서

## 📋 과정 개요
* **과정명:** (창의적이고 전문적인 과정명)
* **교육 대상:** {target}
* **교육 목적:** (개괄식으로 2개 작성)
* **교육 시간:** {total_h}H
* **교육 방식:** (강의형/실습형/워크샵형 중 선택)

---

## 📚 상세 커리큘럼

### 1. [모듈 주제명] (60~90분 제안)
- 주요 학습 내용 요약 (1~2문장)
  - 세부 상세 내용 1
  - 세부 상세 내용 2
  - 세부 상세 내용 3
  - 세부 상세 내용 4
  - 세부 상세 내용 5
  - [실습/토의/롤플레잉] 구체적 활동명 및 실행 방법

(모듈은 교육 시간에 따라 자유롭게 추가. 표 사용 절대 금지)
"""

    try:
        curriculum = _generate(prompt, heavy=True)
    except Exception as e:
        raise Exception(f"A/B 통합 제안서 생성에 실패했습니다: {e}")

    timing_result = validate_curriculum_timing(curriculum, total_h)
    return curriculum, timing_result


# ============ 검수자 AI ============
def review_proposal(proposal_text: str, needs_json: dict) -> dict:
    """
    검수자 페르소나(시니어 HRD 컨설턴트)가 제안서를 평가하고
    점수·피드백·개선지시를 JSON으로 반환
    """
    prompt = f"""당신은 HR 교육 컨설팅 분야의 20년 경력 시니어 컨설턴트이자 제안서 품질 검수 전문가입니다.
아래 AI가 생성한 교육 제안서를 냉정하고 전문적으로 평가하세요.

## 고객 니즈 (기준점)
- 대상: {needs_json.get('target')}
- 핵심 키워드: {', '.join(needs_json.get('core_keywords', []))}
- 문제점: {needs_json.get('pain_point')}
- 기대 행동 변화: {needs_json.get('expected_behavior')}

## 평가 대상 제안서
{proposal_text}

## 평가 기준 (100점 만점)
- 니즈_적합성 (25점): 고객 문제점과 기대 행동 변화가 커리큘럼에 반영되었는가
- 커리큘럼_완성도 (35점): 모듈 흐름(도입→핵심→적용)이 논리적이고 시간 배분이 현실적인가
- 전문성_표현 (25점): HRD 교수법(롤플레잉·진단지·케이스 스터디 등)이 구체적으로 명시되었는가
- 제출_가능성 (15점): 고객사에 바로 이메일로 보낼 수 있는 수준의 문장과 형식인가

반드시 아래 JSON 형식으로만 응답하세요:
{{
  "총점": 85,
  "항목별_점수": {{
    "니즈_적합성": 22,
    "커리큘럼_완성도": 30,
    "전문성_표현": 20,
    "제출_가능성": 13
  }},
  "잘된_점": ["구체적인 강점 1", "강점 2"],
  "개선_필요": ["우선순위 1번 개선사항", "2번", "3번"],
  "즉시_수정_필요": ["고객 제출 전 반드시 고쳐야 할 사항"],
  "제출_가능_여부": "즉시 가능 or 수정 후 가능 or 전면 재작성 필요",
  "개선_지시문": "재생성 AI에게 전달할 구체적인 개선 지시 (2~4문장, 한국어)"
}}"""

    try:
        return json.loads(_generate(prompt, json_mode=True, heavy=True))
    except Exception as e:
        return {"총점": 0, "개선_지시문": str(e), "제출_가능_여부": "오류"}


def improve_proposal(original_proposal: str, review_result: dict,
                     needs_json: dict, duration: int,
                     user_opinion: str = "",
                     retrieved_modules_json: str = None,
                     grouped_modules: dict = None) -> tuple:
    """
    CoT 2단계로 제안서 재작성.
    1단계: 피드백 분석 → 변경 계획 수립 (CoT)
    2단계: 변경 계획 + 모듈 DB → 완전 재작성
    반환: (cot_text, improved_proposal)
    """
    total_h = int(duration)
    total_minutes = total_h * 60

    # 피드백 텍스트 정리
    review_instruction = review_result.get("개선_지시문", "")
    issues   = "\n".join([f"- {i}" for i in review_result.get("개선_필요", [])])
    critical = "\n".join([f"- {i}" for i in review_result.get("즉시_수정_필요", [])])

    if user_opinion.strip():
        feedback_input = f"[사용자 직접 요청]\n{user_opinion.strip()}"
        if review_instruction:
            feedback_input += f"\n\n[AI 검수 결과 — 참고]\n{review_instruction}"
            if issues:
                feedback_input += f"\n개선 필요: {issues}"
            if critical:
                feedback_input += f"\n즉시 수정: {critical}"
    else:
        feedback_input = f"[AI 검수 결과]\n{review_instruction}"
        if issues:
            feedback_input += f"\n개선 필요: {issues}"
        if critical:
            feedback_input += f"\n즉시 수정: {critical}"

    # 모듈 DB 텍스트 정리 (간결하게 — CoT 단계에 과하게 넣지 않음)
    module_names = ""
    if retrieved_modules_json:
        try:
            mods = json.loads(retrieved_modules_json)
            names = [m.get("모듈명", "") for m in mods if m.get("모듈명")]
            module_names = ", ".join(names[:20])
        except Exception:
            module_names = ""

    # ── 1단계: CoT — 변경 계획 수립 ──
    cot_prompt = f"""당신은 시니어 HRD 컨설턴트입니다.
아래 피드백과 기존 제안서를 분석하여, 제안서를 어떻게 개선할지 구체적인 변경 계획을 수립하세요.
계획만 작성하고 제안서는 아직 작성하지 마세요.

## 수정 피드백
{feedback_input}

## 고객 니즈
- 교육 대상: {needs_json.get('target')} | 산업: {needs_json.get('industry')} | {total_h}H
- 핵심 키워드: {', '.join(needs_json.get('core_keywords', []))}
- 문제점: {needs_json.get('pain_point')}
- 기대 행동 변화: {needs_json.get('expected_behavior')}

## 활용 가능한 모듈 목록
{module_names if module_names else "(모듈 목록 없음 — 기존 제안서 내용 기반으로 개선)"}

## 기존 제안서 요약
{original_proposal[:2000]}{"... (이하 생략)" if len(original_proposal) > 2000 else ""}

---

아래 양식으로 변경 계획을 작성하세요:

## 💭 재작성 사고 과정

### 1. 피드백 핵심 파악
- 사용자/검수자가 요청한 핵심 변경 사항은 무엇인가?
- (각 요청을 1줄씩 정리)

### 2. 현재 제안서 문제점 진단
- 기존 제안서에서 피드백 기준으로 부족한 부분은?
- (구체적으로 모듈명·항목 언급)

### 3. 변경 계획
- **삭제/교체할 모듈 또는 내용:** (이유 포함)
- **추가/강화할 모듈 또는 내용:** (이유 포함)
- **구조 변경:** (필요한 경우)

### 4. 재작성 후 기대 결과
- 피드백이 어떻게 반영되었는지 1~2문장으로 예고
"""

    cot_text = ""
    try:
        cot_text = _generate(cot_prompt, heavy=True)
    except Exception as e:
        cot_text = f"(사고 과정 생성 실패: {e})"

    # 모듈 DB 전체 내용 (2단계 재작성용) — 토큰 절약: 내용_원문 400자 제한
    modules_section = ""
    if retrieved_modules_json:
        _slim_for_rewrite = _slim_modules_json(retrieved_modules_json, content_limit=400)
        modules_section = f"""
## 검색된 교육 모듈 (내용_원문에서 활동을 추출해 커리큘럼에 활용)
```json
{_slim_for_rewrite}
```
"""
    elif grouped_modules:
        def _fmt(mods, limit=5):
            out = []
            for m in mods[:limit]:
                meta = m["meta"]
                out.append(f"- [{meta.get('과정명','')}] {meta.get('모듈명','')}: {meta.get('내용_원문','')[:120]}")
            return "\n".join(out) if out else "없음"
        modules_section = f"""
## 교육 모듈 참고
도입: {_fmt(grouped_modules.get('intro',[]))}
핵심: {_fmt(grouped_modules.get('core',[]),8)}
적용: {_fmt(grouped_modules.get('apply',[]))}
"""

    # ── 2단계: 변경 계획 기반 완전 재작성 ──
    rewrite_prompt = f"""당신은 시니어 HRD 컨설턴트입니다.
아래 변경 계획을 100% 실행하여, 교육 제안서를 처음부터 완전히 재작성하세요.

## ⭐ 변경 계획 (반드시 그대로 실행)
{cot_text}

## 수정 피드백 (재확인)
{feedback_input}

## 고객 니즈
- 교육 대상: {needs_json.get('target')} | 산업: {needs_json.get('industry')}
- 핵심 키워드: {', '.join(needs_json.get('core_keywords', []))}
- 문제점: {needs_json.get('pain_point')}
- 기대 행동 변화: {needs_json.get('expected_behavior')}
- 교육 시간: {total_h}H (= {total_minutes}분)
{modules_section}
## 재작성 규칙
- 마크다운 표(table) 절대 금지
- 각 모듈당 세부 항목 5개 이상
- 각 모듈에 [토의]/[실습]/[롤플레잉]/[워크샵]/[사례분석] 중 하나 이상
- 시간 표기: `60~90분 (제안)` 형태
- 내용 축소 절대 금지

반드시 아래 양식으로 작성하세요:

---

# 맞춤 교육 제안서

## 📋 과정 개요
* **과정명:** (창의적이고 전문적인 과정명)
* **교육 대상:** {needs_json.get('target')}
* **교육 목적:** (개괄식으로 2개 작성)
* **교육 시간:** {total_h}H
* **교육 방식:** (강의형/실습형/워크샵형 중 선택)

---

## 📚 상세 커리큘럼

### 1. [모듈 주제명] (60~90분 제안)
- 주요 학습 내용 요약
  - 세부 상세 내용 1
  - 세부 상세 내용 2
  - 세부 상세 내용 3
  - 세부 상세 내용 4
  - 세부 상세 내용 5
  - [실습/토의/롤플레잉] 구체적 활동명 및 실행 방법

(모듈은 교육 시간에 따라 자유롭게 추가. 표 사용 절대 금지)"""

    try:
        improved = _generate(rewrite_prompt, heavy=True)
    except Exception as e:
        raise Exception(f"제안서 재작성 실패: {e}")

    return cot_text, improved



# ============ [P0-1] Placeholder 치환 함수 ============
def replace_placeholders(text: str, company_name: str) -> tuple:
    """
    생성된 텍스트에서 placeholder를 실제 정보로 치환
    Returns: (치환된 텍스트, 미처리된 placeholder 목록)
    """
    result = text
    result = result.replace("[고객사명]", company_name)
    result = result.replace("[고객 사명]", company_name)

    # 미처리된 placeholder 찾기
    remaining = re.findall(r'\[([^\]]+)\]', result)
    remaining = list(set(remaining))  # 중복 제거
    return result, remaining


# ============ Streamlit UI — Editorial Design System ============
st.set_page_config(page_title="제안서 AI", page_icon="📋", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
@import url("https://cdn.jsdelivr.net/gh/orioncactus/pretendard@v1.3.9/dist/web/variable/pretendardvariable.min.css");
@import url("https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;500&display=swap");

:root {
  --canvas:#FAFBFC; --surface:#FFFFFF; --surface-2:#F4F6F8; --tint:#EEF0F3;
  --ink:#0F1419; --ink-2:#4A5568; --ink-3:#94A3B8;
  --line:#ECEEF1; --line-2:#DDE1E6; --line-strong:#C7CCD3;
  --acc:#2A5BFF; --acc-hover:#1A48EA; --acc-soft:rgba(42,91,255,0.08);
  --ok:#15803D; --warn:#B45309; --bad:#B91C1C;
  --ok-bg:#F0FDF4; --warn-bg:#FFFBEB; --bad-bg:#FEF2F2;
  --sans:"Pretendard Variable","Pretendard",-apple-system,BlinkMacSystemFont,"Segoe UI","Apple SD Gothic Neo","Noto Sans KR",system-ui,sans-serif;
  --mono:"JetBrains Mono","SF Mono","Menlo","Consolas",monospace;
  --shadow-card:0 1px 2px rgba(15,23,42,0.04),0 1px 3px rgba(15,23,42,0.06);
  --radius-sm:4px; --radius-md:6px; --radius-lg:10px; --radius-pill:999px;
  --dur-quick:160ms; --ease:cubic-bezier(0.2,0.7,0.2,1);
}

/* ── Streamlit 기본 재정의 ── */
header[data-testid="stHeader"] { display:none !important; }
#MainMenu { display:none !important; }
.stApp { background:var(--canvas) !important; font-family:var(--sans) !important; }
.block-container { padding-top:2rem !important; }
section[data-testid="stSidebar"] { background:var(--surface) !important; border-right:1px solid var(--line) !important; }
section[data-testid="stSidebar"] > div,
section[data-testid="stSidebar"] > div > div { padding-top:0 !important; margin-top:0 !important; }
section[data-testid="stSidebar"] > div { padding-left:16px !important; padding-right:16px !important; padding-bottom:20px !important; }
/* 사이드바 항상 표시 + 닫기 버튼 숨김 */
div[data-testid="stSidebarCollapseButton"] { display:none !important; }
button[data-testid="collapsedControl"] { display:none !important; }
section[data-testid="stSidebar"] {
    display:block !important;
    visibility:visible !important;
    transform:none !important;
    min-width:244px !important;
    margin-left:0 !important;
}
section[data-testid="stSidebar"][aria-expanded="false"] {
    display:block !important;
    transform:translateX(0) !important;
}
.block-container { padding:2rem 3.5rem 6rem !important; max-width:960px !important; }

.stButton > button {
  font-family:var(--sans) !important; font-weight:600 !important; font-size:14px !important;
  border-radius:var(--radius-md) !important; height:38px !important; padding:0 20px !important;
  border:1px solid var(--line-2) !important; background:var(--surface) !important; color:var(--ink) !important;
  transition:all var(--dur-quick) var(--ease) !important;
}
.stButton > button:hover { border-color:var(--ink-2) !important; background:var(--surface-2) !important; }
.stButton > button[kind="primary"] { background:var(--acc) !important; border-color:var(--acc) !important; color:#fff !important; box-shadow:0 1px 2px rgba(42,91,255,0.18) !important; }
.stButton > button[kind="primary"]:hover { background:var(--acc-hover) !important; border-color:var(--acc-hover) !important; }
.stButton > button:disabled { opacity:0.45 !important; cursor:not-allowed !important; }

.stTextArea textarea, .stTextInput input, .stNumberInput input {
  font-family:var(--sans) !important; font-size:15px !important;
  border:1px solid var(--line-2) !important; border-radius:var(--radius-md) !important;
  background:var(--surface) !important; color:var(--ink) !important;
}
.stTextArea textarea:focus, .stTextInput input:focus {
  border-color:var(--acc) !important; box-shadow:0 0 0 3px var(--acc-soft) !important;
}
.stTextArea label,.stTextInput label,.stNumberInput label,.stSelectbox label,.stCheckbox label {
  font-family:var(--sans) !important; font-size:12.5px !important;
  font-weight:600 !important; color:var(--ink) !important;
}
hr { border-color:var(--line) !important; margin:32px 0 !important; }
.stChatMessage { border-radius:var(--radius-md) !important; font-family:var(--sans) !important; }
.stMetric label { font-family:var(--mono) !important; font-size:10.5px !important; text-transform:uppercase !important; letter-spacing:0.08em !important; color:var(--ink-3) !important; }
.stProgress > div { background:var(--tint) !important; border-radius:999px !important; height:5px !important; }
.stProgress > div > div { background:var(--acc) !important; border-radius:999px !important; }
.streamlit-expanderHeader { font-family:var(--mono) !important; font-size:11px !important; letter-spacing:0.07em !important; text-transform:uppercase !important; }

/* ── 커스텀 컴포넌트 ── */
.pai-topbar { display:flex; align-items:center; justify-content:space-between; margin-bottom:36px; padding-bottom:18px; border-bottom:1px solid var(--line); }
.pai-brand { display:inline-flex; align-items:baseline; gap:8px; font-weight:800; font-size:16px; letter-spacing:-0.022em; color:var(--ink); }
.pai-brand .bm { width:22px; height:22px; display:inline-flex; align-items:center; justify-content:center; background:var(--ink); color:#fff; font-style:italic; font-size:15px; line-height:1; border-radius:4px; }
.pai-brand .lt { font-weight:500; color:var(--ink-2); }
.pai-session { font-family:var(--mono); font-size:11px; color:var(--ink-2); letter-spacing:0.04em; display:inline-flex; align-items:center; gap:7px; padding:5px 10px; background:var(--surface-2); border-radius:999px; border:1px solid var(--line); }
.pai-session .dot { width:6px; height:6px; border-radius:50%; background:var(--ok); display:inline-block; }

.pai-section-head { display:grid; grid-template-columns:72px 1fr; align-items:start; gap:20px; margin-bottom:28px; }
.pai-section-num { font-weight:800; font-size:52px; line-height:0.95; letter-spacing:-0.03em; color:var(--acc); font-variant-numeric:tabular-nums; }
.pai-section-title { font-size:26px; line-height:1.2; font-weight:700; letter-spacing:-0.022em; margin:6px 0 6px; color:var(--ink); }
.pai-section-sub { font-size:14px; line-height:1.6; color:var(--ink-2); margin:0; }

.pai-step-collapsed { display:grid; grid-template-columns:52px 1fr; align-items:center; gap:16px; padding:14px 20px; background:var(--surface); border:1px solid var(--line); border-radius:var(--radius-md); margin:0 0 14px; box-shadow:var(--shadow-card); }
.pai-sc-n { font-weight:700; font-size:13px; width:28px; height:28px; display:inline-flex; align-items:center; justify-content:center; background:var(--ink); color:#fff; border-radius:999px; font-variant-numeric:tabular-nums; }
.pai-sc-label { font-family:var(--mono); font-size:10.5px; color:var(--ink-2); letter-spacing:0.08em; text-transform:uppercase; margin-bottom:4px; font-weight:500; }
.pai-sc-summary { font-size:14px; color:var(--ink); line-height:1.45; letter-spacing:-0.005em; }

.pai-summary { border:1px solid var(--line); border-radius:var(--radius-md); background:var(--surface); overflow:hidden; box-shadow:var(--shadow-card); margin-bottom:24px; }
.pai-sum-row { display:grid; grid-template-columns:140px 1fr; align-items:start; gap:20px; padding:14px 20px; border-top:1px solid var(--line); }
.pai-sum-row:first-child { border-top:none; }
.pai-sum-k { font-family:var(--sans); font-size:12.5px; color:var(--ink-2); font-weight:500; }
.pai-sum-v { font-size:15px; line-height:1.45; color:var(--ink); font-weight:500; letter-spacing:-0.011em; }
.pai-sum-v .muted { color:var(--ink-2); font-weight:400; }
.pai-tag-row { display:flex; gap:6px; flex-wrap:wrap; margin-top:2px; }
.pai-tag { font-size:12px; font-weight:500; padding:3px 10px; border:1px solid var(--line-2); border-radius:999px; color:var(--ink); background:var(--surface-2); white-space:nowrap; }

.pai-info-head { display:flex; align-items:baseline; justify-content:space-between; margin:0 0 10px; }
.pai-info-count { font-size:14px; font-weight:600; color:var(--ink); }
.pai-info-count em { font-style:normal; color:var(--acc); font-weight:700; margin-right:4px; }
.pai-info-sort { font-family:var(--mono); font-size:10.5px; color:var(--ink-3); letter-spacing:0.08em; text-transform:uppercase; }
.pai-mod-table { border:1px solid var(--line); border-radius:var(--radius-md); background:var(--surface); overflow:hidden; box-shadow:var(--shadow-card); margin-bottom:8px; }
.pai-mod-thead { display:grid; grid-template-columns:1fr 70px; gap:16px; padding:10px 20px 10px 56px; background:var(--surface-2); border-bottom:1px solid var(--line); font-family:var(--mono); font-size:10px; color:var(--ink-2); letter-spacing:0.1em; text-transform:uppercase; font-weight:500; }
.pai-mod-thead .sc { text-align:right; }
.pai-mod-info { padding:12px 20px 12px 0; }
.pai-mod-rank { font-family:var(--mono); font-weight:600; font-size:12px; color:var(--ink-3); }
.pai-mod-title { font-size:14.5px; font-weight:600; color:var(--ink); letter-spacing:-0.011em; margin-bottom:4px; }
.pai-mod-desc { font-size:12.5px; color:var(--ink-2); display:flex; gap:7px; flex-wrap:wrap; align-items:center; }
.pai-mod-pip { width:3px; height:3px; background:var(--ink-3); border-radius:50%; display:inline-block; }
.pai-mod-sim { font-weight:700; font-size:20px; color:var(--ink); letter-spacing:-0.018em; line-height:1; text-align:right; font-variant-numeric:tabular-nums; }
.pai-mod-sim .pct { font-family:var(--mono); font-size:9.5px; color:var(--ink-3); display:block; margin-top:3px; letter-spacing:0.06em; text-transform:uppercase; }
.pai-ttag { background:#E8F4F8; color:#1A6B9E; padding:2px 9px; border-radius:10px; font-size:0.76rem; border:1px solid #B8D9EE; white-space:nowrap; }
.pai-itag { background:#F0F8E8; color:#2D6A1E; padding:2px 9px; border-radius:10px; font-size:0.76rem; border:1px solid #B8D9B8; white-space:nowrap; }

.pai-modes { display:grid; grid-template-columns:1fr 1fr; gap:12px; margin-top:12px; }
.pai-mode { border:1px solid var(--line-2); border-radius:var(--radius-md); padding:18px 20px; background:var(--surface); position:relative; }
.pai-mode.active { border-color:var(--acc); box-shadow:0 0 0 3px var(--acc-soft); }
.pai-mode .badge { font-family:var(--mono); font-size:11px; color:var(--ink-3); letter-spacing:0.08em; text-transform:uppercase; font-weight:500; margin-right:8px; }
.pai-mode h5 { margin:0 0 6px; font-size:16px; font-weight:700; color:var(--ink); }
.pai-mode p { margin:0; font-size:13px; color:var(--ink-2); line-height:1.55; }
.pai-mode .chk { position:absolute; top:14px; right:16px; font-size:11.5px; color:var(--acc); font-weight:600; }

.pai-meta-bar { display:grid; grid-template-columns:repeat(4,1fr); border:1px solid var(--line); background:var(--surface); border-radius:var(--radius-md); padding:16px 0; margin-bottom:24px; box-shadow:var(--shadow-card); }
.pai-meta-bar > div { padding:0 20px; border-right:1px solid var(--line); }
.pai-meta-bar > div:last-child { border-right:none; }
.pai-meta-k { font-family:var(--mono); font-size:10px; color:var(--ink-2); letter-spacing:0.08em; text-transform:uppercase; margin-bottom:6px; font-weight:500; }
.pai-meta-v { font-size:18px; font-weight:700; letter-spacing:-0.018em; font-variant-numeric:tabular-nums; }
.pai-meta-v .m { color:var(--ink-2); font-weight:400; margin-left:4px; font-size:14px; }

.pai-score-block { display:grid; grid-template-columns:200px 1fr; gap:40px; align-items:center; padding:28px 32px; margin:0 0 16px; background:var(--surface); border:1px solid var(--line); border-radius:var(--radius-md); box-shadow:var(--shadow-card); }
.pai-score-num { font-weight:800; font-size:92px; line-height:0.9; letter-spacing:-0.04em; font-variant-numeric:tabular-nums; }
.pai-score-num.hi { color:var(--ok); } .pai-score-num.md { color:var(--warn); } .pai-score-num.lo { color:var(--bad); }
.pai-score-deno { font-family:var(--mono); font-size:11px; color:var(--ink-2); letter-spacing:0.08em; text-transform:uppercase; margin-top:8px; font-weight:500; }
.pai-score-bars { display:flex; flex-direction:column; gap:14px; }
.pai-score-row { display:grid; grid-template-columns:1fr 55px; gap:10px; align-items:baseline; }
.pai-score-lbl { font-size:13.5px; font-weight:500; color:var(--ink); }
.pai-score-val { font-family:var(--mono); font-size:11px; color:var(--ink-2); text-align:right; font-weight:500; }
.pai-score-track { grid-column:1/-1; height:5px; background:var(--tint); border-radius:999px; overflow:hidden; position:relative; }
.pai-score-fill { position:absolute; top:0; left:0; height:100%; background:var(--acc); border-radius:999px; }
.pai-verdict { padding:14px 20px; background:var(--surface); border:1px solid var(--line); border-radius:var(--radius-md); display:grid; grid-template-columns:140px 1fr; gap:16px; align-items:center; margin-bottom:20px; }
.pai-verdict-k { font-family:var(--mono); font-size:10.5px; color:var(--ink-2); letter-spacing:0.08em; text-transform:uppercase; font-weight:500; }
.pai-verdict-v { font-size:16px; color:var(--ink); font-weight:600; letter-spacing:-0.012em; display:flex; align-items:center; gap:8px; }
.pai-dot { width:8px; height:8px; border-radius:50%; display:inline-block; flex-shrink:0; background:var(--warn); }
.pai-dot.ok { background:var(--ok); } .pai-dot.bad { background:var(--bad); }

.pai-review-cols { display:grid; grid-template-columns:1fr 1fr; gap:16px; margin:16px 0; }
.pai-review-col { border:1px solid var(--line); border-radius:var(--radius-md); background:var(--surface); padding:16px 18px; }
.pai-review-h6 { margin:0 0 12px; font-family:var(--mono); font-size:10.5px; color:var(--ink-2); letter-spacing:0.08em; text-transform:uppercase; font-weight:500; display:flex; align-items:center; gap:8px; }
.pai-review-h6::before { content:""; width:8px; height:8px; border-radius:50%; display:inline-block; flex-shrink:0; }
.pai-review-h6.ok::before { background:var(--ok); } .pai-review-h6.fix::before { background:var(--warn); }
.pai-review-ul { padding:0; margin:0; list-style:none; }
.pai-review-li { padding:8px 0; border-top:1px solid var(--line); font-size:13.5px; line-height:1.5; color:var(--ink); }
.pai-review-li:first-child { border-top:none; padding-top:0; }

.pai-rail-title { font-family:var(--mono); font-size:10.5px; color:var(--ink-3); text-transform:uppercase; letter-spacing:0.1em; margin:0 0 12px; font-weight:500; }
.pai-step-item { display:flex; align-items:center; gap:10px; padding:9px 10px; border-radius:var(--radius-md); margin-bottom:2px; }
.pai-step-item.current { background:var(--acc-soft); }
.pai-step-item.future { opacity:0.5; }
.pai-step-num { width:22px; height:22px; border-radius:999px; display:inline-flex; align-items:center; justify-content:center; font-size:12px; font-weight:600; flex-shrink:0; border:1px solid var(--line-2); background:var(--surface); color:var(--ink-3); font-variant-numeric:tabular-nums; }
.pai-step-item.done .pai-step-num { background:var(--ink); color:#fff; border-color:var(--ink); }
.pai-step-item.current .pai-step-num { background:var(--acc); color:#fff; border-color:var(--acc); }
.pai-step-lbl { font-size:13px; font-weight:500; color:var(--ink-2); letter-spacing:-0.008em; }
.pai-step-item.done .pai-step-lbl { color:var(--ink); }
.pai-step-item.current .pai-step-lbl { color:var(--ink); font-weight:600; }
.pai-rail-foot { margin-top:20px; padding-top:16px; border-top:1px solid var(--line); font-family:var(--mono); font-size:10.5px; color:var(--ink-3); letter-spacing:0.04em; line-height:1.7; }

.pai-divider { height:1px; background:var(--line); margin:40px 0 32px; }

.pai-sub-head { font-size:13px; font-weight:600; color:var(--ink-2); letter-spacing:-0.005em; margin:20px 0 12px; display:flex; align-items:center; gap:8px; padding-left:10px; border-left:3px solid var(--acc); }
</style>
""", unsafe_allow_html=True)

# ── HTML 컴포넌트 헬퍼 ──
def _section_head(num: int, title: str, subtitle: str = "") -> str:
    sub = f'<p class="pai-section-sub">{subtitle}</p>' if subtitle else ""
    return f"""<div class="pai-section-head">
  <div class="pai-section-num">{num:02d}</div>
  <div><h2 class="pai-section-title">{title}</h2>{sub}</div>
</div>"""

def _step_collapsed(num: int, label: str, summary: str) -> str:
    return f"""<div class="pai-step-collapsed">
  <span class="pai-sc-n">{num:02d}</span>
  <div>
    <div class="pai-sc-label">{label}</div>
    <div class="pai-sc-summary">{summary}</div>
  </div>
</div>"""

def _needs_summary_html(nj: dict) -> str:
    kw = "".join(f'<span class="pai-tag">{k}</span>' for k in nj.get("core_keywords", []))
    return f"""<div class="pai-summary">
  <div class="pai-sum-row"><span class="pai-sum-k">Pain Point</span><span class="pai-sum-v">{nj.get("pain_point","")}</span></div>
  <div class="pai-sum-row"><span class="pai-sum-k">핵심 키워드</span><span class="pai-sum-v"><div class="pai-tag-row">{kw}</div></span></div>
  <div class="pai-sum-row"><span class="pai-sum-k">대상</span><span class="pai-sum-v">{nj.get("target","")}</span></div>
  <div class="pai-sum-row"><span class="pai-sum-k">산업</span><span class="pai-sum-v">{nj.get("industry","")}</span></div>
  <div class="pai-sum-row"><span class="pai-sum-k">교육 시간</span><span class="pai-sum-v">{nj.get("duration_hours","")}H <span class="muted">/ {nj.get("preferred_style","")}</span></span></div>
  <div class="pai-sum-row"><span class="pai-sum-k">기대 행동 변화</span><span class="pai-sum-v">{nj.get("expected_behavior", nj.get("expected_outcome",""))}</span></div>
</div>"""

def _score_panel_html(total: int, scores: list, verdict: str) -> str:
    tone = "hi" if total >= 90 else "md" if total >= 70 else "lo"
    bars = "".join(f"""<div class="pai-score-row">
      <span class="pai-score-lbl">{s["label"]}</span>
      <span class="pai-score-val">{s["score"]} / {s["max"]}</span>
      <div class="pai-score-track"><div class="pai-score-fill" style="width:{int(s['score']/s['max']*100)}%"></div></div>
    </div>""" for s in scores)
    dot_cls = "ok" if total >= 90 else ("bad" if total < 60 else "")
    return f"""<div class="pai-score-block">
  <div>
    <div class="pai-score-num {tone}">{total}</div>
    <div class="pai-score-deno">/ 100 점 · 검수자 페르소나</div>
  </div>
  <div class="pai-score-bars">{bars}</div>
</div>
<div class="pai-verdict">
  <span class="pai-verdict-k">제출 가능 여부</span>
  <span class="pai-verdict-v"><span class="pai-dot {dot_cls}"></span>{verdict}</span>
</div>"""

def _review_cols_html(pros: list, cons: list) -> str:
    ph = "".join(f'<li class="pai-review-li">{p}</li>' for p in pros)
    ch = "".join(f'<li class="pai-review-li">{c}</li>' for c in cons)
    return f"""<div class="pai-review-cols">
  <div class="pai-review-col"><div class="pai-review-h6 ok">잘된 점</div><ul class="pai-review-ul">{ph}</ul></div>
  <div class="pai-review-col"><div class="pai-review-h6 fix">개선 필요</div><ul class="pai-review-ul">{ch}</ul></div>
</div>"""

def _rail_html(current: int) -> str:
    labels = ["니즈 입력","니즈 분석 확인","모듈 선택","제안서 생성","AI 검수","재생성"]
    items = ""
    for i, lbl in enumerate(labels):
        n = i + 1
        s = "done" if n < current else "current" if n == current else "future"
        items += f'<div class="pai-step-item {s}"><span class="pai-step-num">{n:02d}</span><span class="pai-step-lbl">{lbl}</span></div>'
    return f'<div class="pai-rail-title">제안 작성 흐름</div>{items}'

def _mode_cards_html(active: str) -> str:
    cards = [
        ("standard","Mode 01","표준 모드","RAG 기반으로 선택된 모듈만 사용해 빠르게 커리큘럼을 조립합니다."),
        ("advanced","Mode 02","고도화 모드","조직 맥락 · Learning Transfer · ROI 평가 섹션이 추가됩니다."),
    ]
    html = '<div class="pai-modes">'
    for key, badge, title, desc in cards:
        cls = "pai-mode active" if active == key else "pai-mode"
        chk = '<span class="chk">✓ 선택됨</span>' if active == key else ""
        html += f'<div class="{cls}">{chk}<h5><span class="badge">{badge}</span>{title}</h5><p>{desc}</p></div>'
    html += "</div>"
    return html

# ── session_state 초기화 ──
_DEFAULTS = {
    "workflow_step": 1,
    "company_name": "",
    "industry": "전산업",
    "target": "팀장/리더급",
    "duration": 8,
    "initial_query": None,
    "chatbot_started": False,        # [Sprint 1-1] 명시적 버튼 클릭 전까지 챗봇 비활성
    "needs_conversation": [],
    "needs_complete": False,
    "needs_json": None,
    "retrieved_modules": [],
    "retrieved_modules_json": None,
    "grouped": None,
    "selected_modules": [],
    "generation_track": "standard",  # [Sprint 2] "standard" | "advanced"
    "advanced_context": {},          # [Sprint 2] 고도화 모드 추가 입력값
    "proposal": None,
    "remaining_placeholders": [],
    "curriculum_timing": None,
    "review": None,
    "improve_cot": None,
    "improved_proposal": None,
    # A/B 초안
    "ab_cot": None,
    "ab_draft_a": None,
    "ab_draft_b": None,
    "ab_feedback": "",
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

current_step = st.session_state.workflow_step

# ── 사이드바 ──
with st.sidebar:
    # ── 로고 위치 조절: margin-top 값을 바꾸면 위아래로 이동 ──
    LOGO_MARGIN_TOP = "0px"   # ← 이 값을 바꿔서 로고 위치 조절
    st.markdown(f"""<div class="pai-brand" style="margin-top:{LOGO_MARGIN_TOP};margin-bottom:20px;">
  <span class="bm">P</span>proposal<span class="lt">.ai</span>
</div>""", unsafe_allow_html=True)
    st.markdown(_rail_html(current_step), unsafe_allow_html=True)
    st.divider()
    _key_idx = st.session_state.get("api_key_index", 0)
    st.markdown(f"""<div class="pai-session" style="margin-bottom:8px;">
  <span class="dot"></span>Gemini 연결됨
</div>""", unsafe_allow_html=True)
    st.caption(f"`{MODEL_LITE}` / `{MODEL_HEAVY}` · 키 {_key_idx + 1}/{len(_API_KEYS)}")
    st.divider()
    if st.button("처음부터 다시 시작", use_container_width=True):
        for _k in list(_DEFAULTS.keys()):
            if _k in st.session_state:
                del st.session_state[_k]
        st.rerun()

# 세션 상태에서 고객 정보 읽기 (Step 1에서 저장됨)
company_name = st.session_state.get("company_name", "")
industry     = st.session_state.get("industry", "전산업")
target       = st.session_state.get("target", "팀장/리더급")
duration     = st.session_state.get("duration", 8)

# ─────────────────────────────────────────────────────────
# STEP 1 : 고객 정보 & 니즈 입력
# ─────────────────────────────────────────────────────────

if current_step == 1:
    if not st.session_state.chatbot_started:
        # ── 입력 화면 (챗봇 시작 전) ──
        st.markdown(_section_head(1, "어떤 교육이 필요한가요?", "자유롭게 적어주세요. AI가 대화하며 교육 대상 · 산업 · 기간 · 실습 비중을 자동으로 수집합니다."), unsafe_allow_html=True)

        initial_query_input = st.text_area(
            "고객의 교육 니즈를 자유롭게 입력하세요",
            height=120,
            placeholder=(
                "예: 우리 회사 신임 팀장들이 MZ세대 팀원들과 소통하는 데 어려움을 겪고 있어요. "
                "특히 면담이나 피드백 상황에서 어떻게 대화해야 할지 몰라 회피하는 경향이 있습니다. "
                "실습 위주로 실제 현장에서 바로 쓸 수 있는 스킬을 익히길 원합니다."
            ),
            key="initial_query_input"
        )

        if initial_query_input and initial_query_input != st.session_state.initial_query:
            st.session_state.initial_query = initial_query_input
            st.session_state.needs_conversation = []
            st.session_state.needs_complete = False

        if initial_query_input:
            if st.button("AI 검수 시작 (니즈 분석 진행) →", type="primary", use_container_width=True, key="start_chatbot"):
                st.session_state.initial_query = initial_query_input
                st.session_state.chatbot_started = True
                st.rerun()

    else:
        # ── 챗봇 진행 중 — 초기 입력은 접어서 표시 ──
        _q = str(st.session_state.initial_query or "")
        st.markdown(
            _step_collapsed(1, "Step 01 · 니즈 입력", _q[:90] + ("…" if len(_q) > 90 else "")),
            unsafe_allow_html=True
        )
        st.markdown('<div class="pai-sub-head">추가 정보 수집</div>', unsafe_allow_html=True)

        full_conv_text = st.session_state.initial_query + " " + " ".join(
            m.get("content", "") for m in st.session_state.needs_conversation if m.get("role") == "user"
        )

        # AI 첫 번째 질문 생성 (대화가 없을 때만)
        if len(st.session_state.needs_conversation) == 0:
            with st.spinner("🤖 AI가 질문을 준비 중..."):
                ai_q = generate_follow_up_questions(
                    st.session_state.initial_query, industry, target, []
                )
            st.session_state.needs_conversation.append({"role": "assistant", "content": ai_q})
            st.rerun()

        # 대화 기록 표시
        for msg in st.session_state.needs_conversation:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])

        # 정보 완성도 표시
        completeness = check_info_completeness(full_conv_text)
        filled = sum([completeness.get("교육인원",False), completeness.get("pain_point",False), completeness.get("실습비중",False)])
        def _chk(v): return "✅" if v else "—"
        st.markdown(f"""<div style="display:flex;gap:8px;margin:12px 0;flex-wrap:wrap;">
          <span class="pai-tag">교육 인원 {_chk(completeness.get('교육인원'))}</span>
          <span class="pai-tag">Pain Point {_chk(completeness.get('pain_point'))}</span>
          <span class="pai-tag">실습 비중 {_chk(completeness.get('실습비중'))}</span>
          <span class="pai-tag" style="color:var(--ink-3)">기존 시도 {_chk(completeness.get('기존시도'))} (선택)</span>
          <span class="pai-tag" style="color:var(--acc);border-color:var(--acc);">수집됨 {filled} / 3</span>
        </div>""", unsafe_allow_html=True)

        # 사용자 추가 답변 입력
        user_chat_in = st.chat_input("답변을 입력하세요...", key="needs_chat_input")
        if user_chat_in:
            st.session_state.needs_conversation.append({"role": "user", "content": user_chat_in})
            new_full = full_conv_text + " " + user_chat_in
            if not is_needs_complete(new_full):
                with st.spinner("🤖 다음 질문 준비 중..."):
                    nq = generate_follow_up_questions(
                        st.session_state.initial_query, industry, target,
                        st.session_state.needs_conversation
                    )
                st.session_state.needs_conversation.append({"role": "assistant", "content": nq})
            st.rerun()

        st.markdown('<div class="pai-divider"></div>', unsafe_allow_html=True)

        needs_ok = is_needs_complete(full_conv_text)
        if not needs_ok:
            st.info("필수 정보(교육인원, Pain Point, 실습비중) 중 최소 2가지 이상 입력하면 다음 단계로 진행할 수 있습니다.")

        if st.button(
            "니즈 분석 진행 →",
            type="primary",
            use_container_width=True,
            disabled=not needs_ok,
            key="step1_next"
        ):
            full_needs_text = st.session_state.initial_query + " " + " ".join(
                m.get("content", "") for m in st.session_state.needs_conversation if m.get("role") == "user"
            )
            with st.spinner("🔍 니즈를 구조화하는 중..."):
                nj = analyze_needs(full_needs_text, industry, target, duration)
            st.session_state.needs_json = nj
            st.session_state.workflow_step = 2
            st.rerun()

else:
    _nj  = st.session_state.get("needs_json") or {}
    _ind = _nj.get("industry") or st.session_state.get("industry", "")
    _tgt = _nj.get("target")   or st.session_state.get("target", "")
    _dur = _nj.get("duration_hours") or st.session_state.get("duration", 8)
    _q   = str(st.session_state.initial_query or "")
    st.markdown(
        _step_collapsed(1, "Step 01 · 니즈 입력", f"{_ind} | {_tgt} | {_dur}H | {_q[:80]}{'…' if len(_q)>80 else ''}"),
        unsafe_allow_html=True
    )

# ─────────────────────────────────────────────────────────
# STEP 2 : 니즈 분석 결과 확인
# ─────────────────────────────────────────────────────────
if current_step >= 2 and st.session_state.needs_json:
    st.markdown('<div class="pai-divider"></div>', unsafe_allow_html=True)
    st.markdown(_section_head(2, "추출된 니즈를 검토하세요.", "AI가 대화에서 추론한 항목입니다. 부정확하면 1단계로 돌아가 수정할 수 있습니다."), unsafe_allow_html=True)

    nj = st.session_state.needs_json
    st.markdown(_needs_summary_html(nj), unsafe_allow_html=True)

    if current_step == 2:
        st.caption("위 분석 결과가 맞으면 다음 단계로 진행하세요.")

        col_ok, col_back = st.columns([3, 1])
        with col_ok:
            if st.button(
                "✅ 분석 결과가 맞습니다 — 모듈 검색으로",
                type="primary",
                use_container_width=True,
                key="step2_confirm"
            ):
                collection, db_type = load_module_db()
                with st.spinner("📚 관련 모듈 검색 중..."):
                    try:
                        r_modules, r_modules_json = search_modules_detailed(
                            collection, nj, db_type
                        )
                        grouped = group_modules_by_type(r_modules, db_type)
                        st.session_state.retrieved_modules = r_modules
                        st.session_state.retrieved_modules_json = r_modules_json
                        st.session_state.grouped = grouped
                        st.session_state.workflow_step = 3
                        st.rerun()
                    except Exception as e:
                        st.error(f"모듈 검색 오류: {e}")
                        st.stop()

        with col_back:
            if st.button("다시 입력", use_container_width=True, key="step2_back"):
                st.session_state.workflow_step = 1
                st.session_state.needs_json = None
                st.rerun()

    else:
        kw_str = ", ".join(nj.get("core_keywords", []))
        st.markdown(
            _step_collapsed(2, "Step 02 · 니즈 분석", f"{nj.get('target','')} · {nj.get('industry','')} · {nj.get('duration_hours','')}H | {kw_str}"),
            unsafe_allow_html=True
        )

# ─────────────────────────────────────────────────────────
# STEP 3 : 모듈 선택
# ─────────────────────────────────────────────────────────
if current_step >= 3 and st.session_state.retrieved_modules:
    st.markdown('<div class="pai-divider"></div>', unsafe_allow_html=True)
    st.markdown(_section_head(3, "커리큘럼에 포함할 모듈을 선택하세요.", f"{len(st.session_state.retrieved_modules)}개 모듈 DB에서 검색된 결과입니다. 선택하지 않으면 AI가 자동으로 조립합니다."), unsafe_allow_html=True)

    r_mods = st.session_state.retrieved_modules

    if current_step == 3:
        col_caption, col_back = st.columns([5, 2])
        with col_caption:
            st.caption("선택하지 않으면 AI가 자동으로 최적 모듈을 조합합니다.")
        with col_back:
            if st.button("← 니즈 다시 입력", use_container_width=True, key="back_to_needs"):
                st.session_state.workflow_step = 1
                st.session_state.retrieved_modules = []
                st.session_state.retrieved_modules_json = None
                st.session_state.ab_draft_a = None
                st.session_state.ab_draft_b = None
                st.session_state.ab_cot = None
                st.rerun()

        # 유사도 → 별점
        def sim_to_stars(s):
            if s >= 90: return "⭐⭐⭐⭐⭐"
            elif s >= 80: return "⭐⭐⭐⭐"
            elif s >= 70: return "⭐⭐⭐"
            elif s >= 60: return "⭐⭐"
            return "⭐"

        type_label = {"intro": "도입", "core": "핵심", "apply": "현업적용"}

        sel_idx_cur = [i for i in range(len(r_mods)) if st.session_state.get(f"mod_sel_{i}", False)]
        st.markdown(f"""<div class="pai-info-head">
          <div class="pai-info-count"><em>{len(sel_idx_cur)}</em> / {len(r_mods)}개 모듈 선택됨</div>
          <span class="pai-info-sort">유사도 기준 정렬</span>
        </div>""", unsafe_allow_html=True)

        for i, mod in enumerate(r_mods):
            mod_name      = mod.get("모듈명", f"모듈 {i+1}")
            course        = mod.get("과정명", "")
            sim           = mod.get("similarity_percent", 0)
            content       = mod.get("내용_원문", "")
            rec_time      = mod.get("권장시간", "")
            target_tags   = mod.get("추천타겟", [])
            industry_tags = mod.get("관련산업", [])
            objective     = mod.get("교육목표", "")
            method        = mod.get("교육방식", "")
            mod_type      = mod.get("모듈성격", "")
            stars         = sim_to_stars(sim)
            type_str      = type_label.get(mod_type, "")

            # 태그 HTML (추천타겟 + 관련산업)
            tag_parts = []
            if type_str:
                tag_parts.append(f'<span class="pai-tag" style="color:var(--ink-2);">{type_str}</span>')
            for t in target_tags:
                tag_parts.append(f'<span class="pai-ttag">{t}</span>')
            for ind in industry_tags:
                tag_parts.append(f'<span class="pai-itag">{ind}</span>')
            if tag_parts:
                st.markdown(
                    f'<div style="display:flex;flex-wrap:wrap;gap:5px;margin-bottom:4px;">{"".join(tag_parts)}</div>',
                    unsafe_allow_html=True
                )

            col_chk, col_info = st.columns([1, 11])
            with col_chk:
                st.checkbox("선택", key=f"mod_sel_{i}", label_visibility="collapsed")
            with col_info:
                meta_parts = []
                if course:   meta_parts.append(course)
                if rec_time: meta_parts.append(rec_time)
                if method:   meta_parts.append(method)
                meta_str = "  ·  ".join(meta_parts)
                expander_label = f"**{i+1:02d} · {mod_name}**  |  {meta_str}  |  {stars} ({sim}%)"
                with st.expander(expander_label):
                    if objective:
                        st.markdown(f"**교육목표:** {objective}")
                    if content:
                        st.markdown("**주요 내용:**")
                        for ln in content.split("\n"):
                            if ln.strip():
                                st.markdown(f"- {ln.strip()}")

            st.markdown('<div style="height:1px;background:var(--line);margin:2px 0 6px;"></div>', unsafe_allow_html=True)

        # 선택 현황 요약
        sel_idx = [i for i in range(len(r_mods)) if st.session_state.get(f"mod_sel_{i}", False)]
        if len(sel_idx) == 0:
            st.info("💡 모듈을 선택하지 않으면 AI가 자동으로 선택합니다.")
        else:
            sel_names = [r_mods[i].get("모듈명", f"모듈 {i+1}") for i in sel_idx]
            st.success(f"✅ {len(sel_idx)}개 선택됨: {', '.join(sel_names)}")

        st.markdown('<div class="pai-divider"></div>', unsafe_allow_html=True)

        # ── 생성 방식 선택 ──
        st.markdown('<div class="pai-sub-head">제안서 생성 방식</div>', unsafe_allow_html=True)
        col_std, col_adv = st.columns(2)
        with col_std:
            std_selected = st.session_state.generation_track == "standard"
            if st.button(
                "📄 표준 모드  ·  빠른 뼈대 제안서",
                use_container_width=True,
                type="primary" if std_selected else "secondary",
                key="track_standard"
            ):
                st.session_state.generation_track = "standard"
                st.rerun()
        with col_adv:
            adv_selected = st.session_state.generation_track == "advanced"
            if st.button(
                "🔬 고도화 모드  ·  심층 맞춤 제안서",
                use_container_width=True,
                type="primary" if adv_selected else "secondary",
                key="track_advanced"
            ):
                st.session_state.generation_track = "advanced"
                st.rerun()

        if st.session_state.generation_track == "standard":
            st.caption("표준 모드: RAG 기반 커리큘럼을 빠르게 생성합니다.")
        else:
            st.caption("고도화 모드: 조직 맥락 + Learning Transfer + ROI 평가 섹션이 추가됩니다.")

        # 고도화 모드 추가 입력
        if st.session_state.generation_track == "advanced":
            with st.expander("고도화 추가 정보 입력 (선택)", expanded=True):
                adv_status   = st.text_area("교육 대상 현황", placeholder="예: 팀장 승진 후 평균 6개월 이내, 실무자 출신으로 리더십 경험 부족", height=80, key="adv_status")
                adv_culture  = st.text_area("조직 문화 / 특이사항", placeholder="예: 수평적 문화 지향, 최근 조직 개편으로 팀 갈등 높음", height=80, key="adv_culture")
                adv_change   = st.text_area("기대하는 변화 (구체적)", placeholder="예: 교육 3개월 후 팀원 만족도 10% 향상, 이직률 감소", height=80, key="adv_change")
                adv_special  = st.text_area("특별 요청 사항", placeholder="예: 경쟁사 사례 배제, 특정 강사 선호, 외부 강사 불가", height=80, key="adv_special")
                st.session_state.advanced_context = {
                    "교육대상현황": adv_status,
                    "조직문화": adv_culture,
                    "기대변화": adv_change,
                    "특별요청": adv_special,
                }

        st.divider()

        # ── A/B 초안 생성 버튼 ──
        if not st.session_state.ab_draft_a:
            if st.button("🧠 A/B 커리큘럼 초안 생성", type="primary", use_container_width=True, key="step3_ab"):
                final_sel_idx = [i for i in range(len(r_mods)) if st.session_state.get(f"mod_sel_{i}", False)]
                st.session_state.selected_modules = final_sel_idx
                sel_details = [r_mods[i] for i in final_sel_idx]
                with st.spinner("💭 AI가 두 가지 커리큘럼 방향을 설계하는 중..."):
                    try:
                        cot, draft_a, draft_b = assemble_curriculum_ab(
                            st.session_state.needs_json,
                            st.session_state.retrieved_modules_json,
                            duration,
                            sel_details if sel_details else None,
                        )
                        st.session_state.ab_cot     = cot
                        st.session_state.ab_draft_a = draft_a
                        st.session_state.ab_draft_b = draft_b
                        st.rerun()
                    except Exception as e:
                        st.warning(f"⚠️ A/B 초안 생성 중 오류가 발생했습니다. 잠시 후 다시 시도해 주세요.\n\n{e}")

        # ── A/B 결과 UI ──
        if st.session_state.ab_draft_a:
            if st.session_state.ab_cot:
                with st.expander("💭 AI 설계 사고 과정", expanded=False):
                    st.markdown(st.session_state.ab_cot)

            st.markdown("#### A안 · B안 완성 제안서")
            tab_a, tab_b = st.tabs(["📘 A안", "📗 B안"])
            with tab_a:
                st.markdown(st.session_state.ab_draft_a)
            with tab_b:
                st.markdown(st.session_state.ab_draft_b)

            ab_feedback = st.text_area(
                "✏️ 의견 작성 (선택사항 — AI가 참고하여 A·B안을 통합)",
                placeholder="예: A안의 도입 방식과 B안의 실습 구성을 결합해주세요. 실습 비중을 늘려주세요.",
                height=80,
                key="ab_feedback_input",
            )
            st.session_state.ab_feedback = ab_feedback

            col_reset, col_gen = st.columns([1, 3])
            with col_reset:
                if st.button("↩️ 다시 생성", use_container_width=True, key="ab_reset"):
                    st.session_state.ab_draft_a = None
                    st.session_state.ab_draft_b = None
                    st.session_state.ab_cot     = None
                    st.rerun()
            with col_gen:
                if st.button("🔀 A·B안 통합 최선 제안서 생성", type="primary", use_container_width=True, key="step3_generate"):
                    with st.spinner("A·B안을 비교하고 최선 요소를 통합하는 중..."):
                        proposal, timing_result = combine_ab_proposals(
                            st.session_state.ab_draft_a,
                            st.session_state.ab_draft_b,
                            st.session_state.ab_feedback,
                            st.session_state.needs_json,
                            duration,
                        )

                    proposal = re.sub(r'<br\s*/?>', '\n- ', proposal)
                    proposal = re.sub(r'<[^>]+>', '', proposal)
                    proposal, rem_ph = replace_placeholders(proposal, company_name)

                    st.session_state.proposal = proposal
                    st.session_state.curriculum_timing = timing_result
                    st.session_state.remaining_placeholders = rem_ph
                    st.session_state.workflow_step = 4
                    st.rerun()

    else:
        sel_idx_done = st.session_state.selected_modules
        track_label = "표준" if st.session_state.generation_track == "standard" else "고도화"
        if not sel_idx_done:
            st.markdown(_step_collapsed(3, "Step 03 · 모듈 선택", f"AI 자동 선택 · {track_label} 모드"), unsafe_allow_html=True)
        else:
            done_names = [r_mods[i].get("모듈명", f"모듈 {i+1}") for i in sel_idx_done if i < len(r_mods)]
            preview = ", ".join(done_names[:3]) + ("…" if len(done_names) > 3 else "")
            st.markdown(_step_collapsed(3, "Step 03 · 모듈 선택", f"{len(sel_idx_done)}개 모듈 · {track_label} 모드 | {preview}"), unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────
# STEP 4 : 생성된 제안서
# ─────────────────────────────────────────────────────────
if current_step >= 4 and st.session_state.proposal:
    st.markdown('<div class="pai-divider"></div>', unsafe_allow_html=True)
    st.markdown(_section_head(4, "맞춤 제안서가 완성됐습니다.", "A안 · B안 초안을 비교한 뒤 최적 통합본을 확인하세요."), unsafe_allow_html=True)

    proposal      = st.session_state.proposal
    nj            = st.session_state.needs_json or {}
    timing_result = st.session_state.curriculum_timing
    rem_ph        = st.session_state.remaining_placeholders or []
    keyword       = nj.get("core_keywords", ["제안서"])[0]

    # 제안서 메타 바
    _sel_cnt = len(st.session_state.selected_modules) or len(st.session_state.retrieved_modules)
    _tgt  = nj.get("target", "")
    _dur  = nj.get("duration_hours", duration)
    _sty  = nj.get("preferred_style", "혼합형")
    _trk  = "표준" if st.session_state.generation_track == "standard" else "고도화"
    st.markdown(f"""<div class="pai-meta-bar">
      <div><div class="pai-meta-k">대상</div><div class="pai-meta-v">{_tgt}</div></div>
      <div><div class="pai-meta-k">기간</div><div class="pai-meta-v">{_dur}H</div></div>
      <div><div class="pai-meta-k">방식</div><div class="pai-meta-v">{_sty} <span class="m">· {_trk}</span></div></div>
      <div><div class="pai-meta-k">사용 모듈</div><div class="pai-meta-v">{_sel_cnt}<span class="m"> 개</span></div></div>
    </div>""", unsafe_allow_html=True)

    st.markdown(proposal)
    st.markdown('<div class="pai-divider"></div>', unsafe_allow_html=True)

    st.markdown('<div style="font-size:12.5px;font-weight:600;color:var(--ink);margin-bottom:8px;">제출용 버전 다운로드</div>', unsafe_allow_html=True)
    try:
        docx_bytes = markdown_to_docx(proposal)
        docx_ok = True
    except Exception:
        docx_ok = False

    dl_c1, dl_c2 = st.columns(2)
    with dl_c1:
        if docx_ok:
            st.download_button(
                "📄 Word 다운로드 (.docx)",
                data=docx_bytes,
                file_name=f"제안서_{keyword}_제출용.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )
    with dl_c2:
        st.download_button(
            "📥 Markdown (.md)",
            data=proposal,
            file_name=f"제안서_{keyword}.md",
            mime="text/markdown",
            use_container_width=True
        )

    # ── [P2] 내부 QA 버전 (expander — 제출용 파일에는 미포함) ──
    with st.expander("🔧 내부 검토용 (QA 버전)", expanded=False):

        # [P1] Self-Correction 품질 검증 결과 표시
        st.markdown("### 🤖 Self-Correction 품질 검증")
        quality_check = validate_proposal_quality(proposal)
        if quality_check["passed"]:
            st.success(
                f"✅ 품질 기준 통과 | "
                f"모듈 {quality_check['module_count']}개 | "
                f"상호작용: {', '.join(quality_check['interactions_found']) or '없음'}"
            )
        else:
            st.warning(
                "⚠️ 최종 품질 체크 (참고용)\n" +
                "\n".join(f"- {f}" for f in quality_check["failures"])
            )

        st.markdown("### 📊 시간 검증 결과 (참고용)")
        if timing_result:
            tc1, tc2 = st.columns([1, 2])
            with tc1:
                alloc_m = timing_result["total_minutes_allocated"]
                st.metric("배분된 교육 시간 (참고)", f"{alloc_m}분 ({alloc_m/60:.1f}H)")
            with tc2:
                if timing_result["valid"]:
                    st.success(f"✅ 시간 범위 내 | 여유: {timing_result['variance_minutes']}분")
                else:
                    st.info(
                        f"ℹ️ 시간 초과 {abs(timing_result['variance_minutes'])}분 — "
                        "풍부한 내용 제공 목적이므로 컨설턴트가 현장에서 조율하세요."
                    )
            # 시간 초과 경고는 오류가 아닌 참고 정보로 표시
            for w in timing_result.get("warnings", []):
                st.caption(f"📌 참고: {w}")

        st.markdown("### 🔍 Placeholder 검증")
        if rem_ph:
            st.warning(f"⚠️ 미처리 항목: {', '.join(f'[{p}]' for p in rem_ph)}\n수동으로 입력해주세요.")
        else:
            st.success("✅ 모든 placeholder 처리 완료")

        st.markdown("### 📋 모듈 사용 정보")
        sel_cnt = len(st.session_state.selected_modules)
        if sel_cnt:
            st.info(f"사용자 선택 모듈 {sel_cnt}개 우선 포함")
        else:
            st.info("AI 자동 선택 모드")
        for m in st.session_state.retrieved_modules[:5]:
            st.markdown(f"- {m.get('모듈명', '')} ({m.get('similarity_percent', 0)}%)")

        # QA 버전 다운로드
        qa_md = (
            f"# QA 검토 버전\n\n{proposal}\n\n---\n\n## 내부 검토 정보\n"
        )
        if timing_result:
            qa_md += f"\n### 시간 검증\n{timing_result['details']}\n"
        if rem_ph:
            qa_md += f"\n### 미처리 항목\n{', '.join(rem_ph)}\n"
        st.download_button(
            "📥 내부 검토용 다운로드 (.md)",
            data=qa_md,
            file_name=f"제안서_{keyword}_QA버전.md",
            mime="text/markdown",
            use_container_width=True
        )

    # ── STEP 5 : AI 검수 ──
    st.markdown('<div class="pai-divider"></div>', unsafe_allow_html=True)
    st.markdown(_section_head(5, "검수자 페르소나의 평가입니다.", "시니어 HRD 컨설턴트 관점에서 4가지 기준으로 점수화했습니다."), unsafe_allow_html=True)

    if st.button("AI 검수 시작 →", type="primary", use_container_width=False, key="review_btn"):
        with st.spinner("시니어 HRD 컨설턴트가 제안서를 검토 중..."):
            review = review_proposal(
                st.session_state.improved_proposal or proposal,
                nj
            )
            st.session_state.review = review

    if st.session_state.review:
        review = st.session_state.review

        if review.get("제출_가능_여부") == "오류":
            st.error(f"AI 검수 중 오류가 발생했습니다: {review.get('개선_지시문', '')}")
            st.caption("'AI 검수 시작' 버튼을 다시 눌러 재시도하거나, 아래 6단계에서 직접 피드백을 입력해 재작성할 수 있습니다.")
        else:
            total   = review.get("총점", 0)
            verdict = review.get("제출_가능_여부", "")
            scores_raw  = review.get("항목별_점수", {})
            max_map = {"니즈_적합성": 25, "커리큘럼_완성도": 35, "전문성_표현": 25, "제출_가능성": 15}
            score_list = [{"label": k, "score": v, "max": max_map.get(k, 25)} for k, v in scores_raw.items()]

            st.markdown(_score_panel_html(total, score_list, verdict), unsafe_allow_html=True)
            st.markdown(_review_cols_html(review.get("잘된_점", []), review.get("개선_필요", [])), unsafe_allow_html=True)

            if review.get("즉시_수정_필요"):
                st.warning("즉시 수정 필수\n" + "\n".join(f"- {i}" for i in review["즉시_수정_필요"]))

            st.info(f"검수자 개선 지시문:\n{review.get('개선_지시문', '')}")

    # ── STEP 6 : 피드백 반영 재생성 ──
    st.markdown('<div class="pai-divider"></div>', unsafe_allow_html=True)
    st.markdown(_section_head(6, "피드백을 반영해 재생성하세요.", "검수 의견을 그대로 반영하거나, 추가 수정 요청을 자유롭게 입력할 수 있습니다."), unsafe_allow_html=True)

    user_opinion = st.text_area(
        "추가 수정 요청 — 선택",
        placeholder="예: 3모듈의 실습 시간을 늘려주세요. 사례 연구를 금융권 중심으로 바꿔주세요.\n비워두면 위 AI 검수 결과를 기반으로 재작성합니다.",
        height=100,
        key="user_opinion_input"
    )

    _has_valid_review = (
        st.session_state.review is not None
        and st.session_state.review.get("제출_가능_여부") != "오류"
    )
    _can_rewrite = bool(user_opinion.strip()) or _has_valid_review

    if not _can_rewrite:
        st.caption("수정 요청을 입력하거나, 먼저 'AI 검수 시작'을 실행하세요.")

    if st.button(
        "🔄 제안서 재작성",
        use_container_width=True,
        disabled=not _can_rewrite,
        key="improve_btn"
    ):
        _review_for_improve = (
            st.session_state.review if _has_valid_review
            else {"개선_지시문": "", "개선_필요": [], "즉시_수정_필요": []}
        )
        with st.spinner("피드백을 분석하고 제안서를 재작성하는 중..."):
            cot_result, improved = improve_proposal(
                st.session_state.improved_proposal or proposal,
                _review_for_improve,
                nj,
                duration,
                user_opinion=user_opinion,
                retrieved_modules_json=st.session_state.retrieved_modules_json,
                grouped_modules=st.session_state.grouped,
            )
        improved = re.sub(r'<br\s*/?>', '\n- ', improved)
        improved = re.sub(r'<[^>]+>', '', improved)
        improved, _ = replace_placeholders(improved, company_name)
        st.session_state.improve_cot = cot_result
        st.session_state.improved_proposal = improved

    if st.session_state.improved_proposal:
        if st.session_state.improve_cot:
            with st.expander("💭 AI 재작성 사고 과정", expanded=False):
                st.markdown(st.session_state.improve_cot)
        st.success("✅ 재작성된 제안서입니다.")
        st.markdown(st.session_state.improved_proposal)
        st.divider()

        imp_timing = validate_curriculum_timing(
            st.session_state.improved_proposal,
            int(duration)
        )
        if imp_timing:
            itc1, itc2 = st.columns([1, 2])
            with itc1:
                im = imp_timing["total_minutes_allocated"]
                st.metric("재작성본 배분 시간", f"{im}분 ({im/60:.1f}H)")
            with itc2:
                if imp_timing["valid"]:
                    st.success(f"✅ 시간 정합성 OK | 여유: {imp_timing['variance_minutes']}분")
                else:
                    st.error(f"❌ 시간 초과: {abs(imp_timing['variance_minutes'])}분 과다")

        try:
            imp_docx = markdown_to_docx(st.session_state.improved_proposal)
            st.download_button(
                "📄 재작성본 Word 다운로드 (.docx)",
                data=imp_docx,
                file_name=f"제안서_{keyword}_재작성본.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )
        except Exception:
            pass

        st.caption("'AI 검수 시작' 버튼을 다시 누르면 재작성본을 검수합니다.")

st.markdown('<div class="pai-divider"></div>', unsafe_allow_html=True)
st.markdown('<div style="font-family:var(--mono);font-size:10.5px;color:var(--ink-3);letter-spacing:0.06em;">Powered by KPH · proposal.ai</div>', unsafe_allow_html=True)
